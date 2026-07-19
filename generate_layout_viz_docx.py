"""
FBSL-KAGS Layout Generation & Visualization Pipeline
Word Document Generator
Covers: Tech stack, FBSL->image pipeline, adjacency graph, SVG floor plan
"""
import sys, traceback
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn
except Exception as e:
    print(f"IMPORT ERROR: {e}", file=sys.stderr); sys.exit(1)


# ── helpers ────────────────────────────────────────────────────────────────────
def shade(cell, color):
    tc = cell._element.get_or_add_tcPr()
    tc.append(tc.makeelement(qn('w:shd'), {qn('w:fill'): color, qn('w:val'): 'clear'}))

def hdr(table, cells, bg='1F3864'):
    row = table.add_row()
    for i, t in enumerate(cells):
        c = row.cells[i]; p = c.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT if i == 0 else WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(str(t)); r.font.size = Pt(9); r.font.name = 'Calibri'
        r.bold = True; r.font.color.rgb = RGBColor(255, 255, 255); shade(c, bg)

def drow(table, cells, bold=False, bg=None):
    row = table.add_row()
    for i, t in enumerate(cells):
        c = row.cells[i]; p = c.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT if i == 0 else WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(str(t)); r.font.size = Pt(9); r.font.name = 'Calibri'
        if bold: r.bold = True
        if bg: shade(c, bg)

def mk(doc, n_cols, headers, rows, bold_first=False, bg='1F3864'):
    t = doc.add_table(rows=1, cols=n_cols)
    t.style = 'Table Grid'; t.alignment = WD_TABLE_ALIGNMENT.LEFT
    hdr(t, headers, bg=bg); t._element.remove(t.rows[0]._element)
    for i, row in enumerate(rows):
        drow(t, row, bold=(bold_first and i == 0))
    return t

def h1(doc, text):
    h = doc.add_heading(text, level=1)
    for r in h.runs: r.font.color.rgb = RGBColor(0x1F, 0x38, 0x64)

def h2(doc, text): return doc.add_heading(text, level=2)
def h3(doc, text): return doc.add_heading(text, level=3)

def para(doc, text, bold=False, italic=False):
    p = doc.add_paragraph(); r = p.add_run(text)
    r.font.size = Pt(11); r.font.name = 'Calibri'; r.bold = bold; r.italic = italic
    return p

def code(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.35)
    r = p.add_run(text)
    r.font.name = 'Courier New'; r.font.size = Pt(9)
    r.font.color.rgb = RGBColor(0x1F, 0x38, 0x64)

def bul(doc, label, rest):
    p = doc.add_paragraph(style='List Bullet')
    if label:
        r1 = p.add_run(label); r1.bold = True
        r1.font.size = Pt(10); r1.font.name = 'Calibri'
    r2 = p.add_run(rest); r2.font.size = Pt(10); r2.font.name = 'Calibri'

def note(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.25)
    r = p.add_run(text); r.italic = True
    r.font.size = Pt(10); r.font.name = 'Calibri'
    r.font.color.rgb = RGBColor(0x59, 0x56, 0x59)


# ── build ──────────────────────────────────────────────────────────────────────
def build():
    doc = Document()
    ns = doc.styles['Normal']
    ns.font.name = 'Calibri'; ns.font.size = Pt(11)
    ns.paragraph_format.space_after = Pt(6)
    ns.paragraph_format.line_spacing = 1.15

    # ── TITLE ─────────────────────────────────────────────────────────────────
    t = doc.add_heading('FBSL-KAGS Layout Generation & Visualization', level=0)
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for r in t.runs: r.font.size = Pt(22); r.font.color.rgb = RGBColor(0x1F, 0x38, 0x64)

    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run('How the Pipeline Converts FBSL Nodes into Floor Plan Images and Adjacency Graphs')
    r.font.size = Pt(13); r.bold = True; r.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)

    p2 = doc.add_paragraph(); p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run('Tech Stack  |  SVG Floor Plan  |  Adjacency Graph  |  Image Rendering')
    r2.font.size = Pt(11); r2.italic = True; r2.font.color.rgb = RGBColor(0x59, 0x56, 0x59)

    doc.add_paragraph('')

    # ── 1. TECH STACK ─────────────────────────────────────────────────────────
    h1(doc, '1.  Full Technology Stack')
    para(doc,
        'The layout generation and visualization pipeline uses seven distinct libraries. '
        'Each serves a specific role — no library is interchangeable with another.')

    mk(doc, 4,
        ['Library', 'Version Role', 'Used For', 'Source File'],
        [
            ('Shapely',         'Geometry engine',       'Room polygon creation, adjacency detection (.touches(), .distance()), bounding boxes, Shapely.box()', 'layout_agent.py, enhanced_layout.py'),
            ('NumPy',           'Numerical arrays',      'Force-directed position arrays, distance calculations, matrix operations (adjacency_matrix)', 'spatial_algorithms.py'),
            ('scipy.optimize',  'Constrained solver',    'SLSQP post-optimisation (minimize() with equality constraints on room area)', 'layout_agent.py'),
            ('NetworkX',        'Graph data structure',  'Adjacency graph nodes/edges, spring_layout(), connected_components(), is_connected()', 'layout_agent.py, enhanced_layout.py'),
            ('Matplotlib',      'Raster image output',   '18x9 inch dual-panel figure (spatial layout + adjacency graph) -> PNG at 200 DPI', 'enhanced_layout.py'),
            ('xml.etree.ElementTree', 'SVG construction', 'Building SVG DOM tree (rooms, labels, circulation, legend, compass, grid)', 'visualization.py, enhanced_layout.py'),
            ('heapq',           'Priority queue',        'A* open-set heap (f-score ordering) for circulation path finding', 'spatial_algorithms.py'),
            ('FAISS (faiss-cpu)','Vector search',        'RAG retrieval of Finnish floor plan precedents; IndexFlatIP cosine similarity', 'research_agent.py'),
        ])

    doc.add_paragraph('')

    # ── 2. END-TO-END PIPELINE ────────────────────────────────────────────────
    h1(doc, '2.  End-to-End: FBSL Node to Visual Output')
    para(doc,
        'The journey from an abstract FBSL node (containing Functions, Behaviors, Structures, '
        'Layout rooms) to a floor plan image and adjacency graph PNG passes through 5 stages:')

    mk(doc, 4,
        ['Stage', 'Input', 'Output', 'Code Location'],
        [
            ('1. Room Specs',     'FBSLLayoutNode.layout.rooms (Room objects with .area)',       '12 room_specs dicts {area, width, length, min_w, max_w}', '_calculate_room_specifications()  layout_agent.py:397'),
            ('2. Adjacency Matrix','Room objects with .required_adjacencies, .preferred_adjacencies, .avoid_adjacencies', '12x12 NumPy ndarray w[i,j] in [-1, +1]', '_build_adjacency_matrix()  layout_agent.py:436'),
            ('3. Positions',      '12 room_specs + 12x12 adj matrix',                          'Dict {room_id: {x, y, width, height}} in metre coordinates', '_force_directed_placement() -> _optimize_layout()  layout_agent.py:460'),
            ('4. Polygons',       '{room_id: {x, y, width, height}}',                          '{room_id: Shapely Polygon}  12 rectangular polygons', '_generate_room_polygons()  layout_agent.py:623'),
            ('5a. SVG Floor Plan','Shapely polygons + circulation paths',                       'SVG string (XML): rooms, labels, dashed circulation, legend, metrics overlay', 'SVGFloorPlanGenerator.generate_floor_plan()  utils/visualization.py:82'),
            ('5b. Enhanced SVG', 'Shapely polygons via LayoutRoomAdapter',                      'SVG file saved to visual_outputs/ with compass rose, 1m grid, colour-coded rooms', 'EnhancedLayoutVisualizer.render()  visualization/enhanced_layout.py:222'),
            ('5c. Adj PNG',       'NetworkX graph with edge types',                             '18x9 in PNG at 200 DPI: spatial floor plot (left) + spring-layout graph (right)', '_generate_adjacency_graph()  enhanced_layout.py:388'),
        ])

    doc.add_paragraph('')

    # ── 3. STAGE 1: ROOM SPECS ────────────────────────────────────────────────
    h1(doc, '3.  Stage 1 — Room Specification Extraction')
    para(doc,
        'For each Room in node.layout.rooms, the agent computes a square approximation '
        'of its area. This initial square is then relaxed by the SLSQP solver in Stage 3.')

    code(doc,
        'for room_id, room in node.layout.rooms.items():\n'
        '    area = room.area  if room.area > 0  else 16.0  # default 16 m2\n'
        '    side = np.sqrt(area)                            # square root approximation\n'
        '\n'
        '    specs[room_id] = {\n'
        '        "area"     : area,\n'
        '        "width"    : side,\n'
        '        "length"   : side,\n'
        '        "height"   : room.height or 3.0,\n'
        '        "min_width": max(2.0, side * 0.7),   # min_room_dimension = 2.0 m\n'
        '        "max_width": min(15.0, side * 1.3),  # max_room_dimension = 15.0 m\n'
        '    }')

    para(doc, 'For the 4-bedroom family home:')
    mk(doc, 4,
        ['Room', 'Area (m2)', 'Initial Side (m)', 'Allowed Width Range (m)'],
        [
            ('Master Bedroom',   '21.0', '4.58', '3.21 - 5.96'),
            ('Master Bathroom',  '7.0',  '2.65', '2.00 - 3.44'),
            ('Child Bedroom x3', '14.0', '3.74', '2.62 - 4.86'),
            ('Living/Dining',    '40.0', '6.32', '4.43 - 8.22'),
            ('Kitchen',          '16.0', '4.00', '2.80 - 5.20'),
            ('Home Office',      '12.0', '3.46', '2.42 - 4.50'),
            ('Laundry/Mudroom/Storage', '5.0', '2.24', '2.00 - 2.91'),
            ('Shared Bathroom',  '6.0',  '2.45', '2.00 - 3.18'),
        ])

    doc.add_paragraph('')

    # ── 4. STAGE 2: ADJACENCY MATRIX ─────────────────────────────────────────
    h1(doc, '4.  Stage 2 — Weighted Adjacency Matrix  (WeightedAdjacencyCalculator)')
    para(doc,
        'WeightedAdjacencyCalculator builds the 12x12 matrix that drives both the '
        'force simulation (Stage 3) and the A* path selection (Stage 4). '
        'Three relationship types are extracted from room metadata:')

    code(doc,
        '# extract_adjacency_from_fbsl()  (spatial_algorithms.py:471)\n'
        'for room in rooms:\n'
        '    for adj_id in room.required_adjacencies:\n'
        '        functional_deps[(room_id, adj_id)] = 1.0   # strong functional link\n'
        '        traffic_flows[(room_id, adj_id)]   = 0.8   # high foot traffic\n'
        '\n'
        '    for adj_id in room.preferred_adjacencies:\n'
        '        functional_deps[(room_id, adj_id)] = 0.6   # moderate link\n'
        '        traffic_flows[(room_id, adj_id)]   = 0.5\n'
        '\n'
        '    for adj_id in room.avoid_adjacencies:\n'
        '        privacy_reqs[(room_id, adj_id)]    = 1.0   # keep apart\n'
        '\n'
        '# calculate_adjacency_matrix()  (spatial_algorithms.py:394)\n'
        'A = alpha * func_matrix + beta * traffic_matrix + gamma * (-privacy_matrix)\n'
        '  where alpha=0.40, beta=0.35, gamma=0.25  (normalised from 0.40/0.35/0.25)\n'
        '\n'
        'A = A / abs(A).max()   # normalise to [-1, +1]')

    para(doc, 'Resulting matrix values for key room pairs:')
    mk(doc, 5,
        ['Room Pair', 'FD', 'TF', 'Privacy', 'Final w[i,j]'],
        [
            ('Kitchen <-> Living/Dining',     '1.0', '0.8', '0.0', '+0.714  (strongly attract)'),
            ('Master Bedroom <-> Bathroom',   '1.0', '0.8', '0.0', '+0.714'),
            ('Child Bedroom <-> Shared Bath', '0.6', '0.5', '0.0', '+0.435'),
            ('Mudroom <-> Laundry',           '0.6', '0.5', '0.0', '+0.435'),
            ('Master Bedroom <-> Kitchen',    '0.0', '0.0', '1.0', '-0.263  (repel)'),
            ('Storage <-> Living/Dining',     '0.0', '0.0', '0.0', '0.000  (neutral)'),
        ])

    note(doc,
        'The matrix is symmetric. Positive = rooms should be placed adjacent. '
        'Negative = rooms should be separated. Zero = no spatial preference.')

    doc.add_paragraph('')

    # ── 5. STAGE 3: FORCE-DIRECTED ────────────────────────────────────────────
    h1(doc, '5.  Stage 3 — Force-Directed Placement + SLSQP Refinement')

    h2(doc, '5.1  Phase A: Force-Directed Simulation  (ForceDirectedLayout)')
    para(doc,
        'Rooms start at random positions within a 50x50m grid. The physics engine '
        'iterates up to 200 times until max_displacement < 0.01m:')

    code(doc,
        '# spatial_algorithms.py:585\n'
        'for iteration in range(200):\n'
        '    for each pair (r_i, r_j):\n'
        '        d  = distance(r_i, r_j) + 0.01      # epsilon avoids div/0\n'
        '        u  = direction_unit_vector(r_i, r_j)\n'
        '\n'
        '        # Attraction (only when adjacency weight > 0)\n'
        '        if A[i,j] > 0:\n'
        '            F_attract = k_attr * A[i,j] * d * u   # k_attr = 0.1\n'
        '            forces[r_i] += F_attract\n'
        '        elif A[i,j] < 0:\n'
        '            F_repel = k_attr * |A[i,j]| * d * u\n'
        '            forces[r_i] -= F_repel             # push away\n'
        '\n'
        '        # Universal repulsion (ALL pairs)\n'
        '        F_repulse = (k_rep / d^2) * u          # k_rep = 100.0\n'
        '        forces[r_i] -= F_repulse\n'
        '\n'
        '    for each room r_i:\n'
        '        p_i += learning_rate * forces[r_i]     # lr = 0.01\n'
        '        p_i = clip(p_i, 0, 100)               # stay in bounds\n'
        '        max_displacement = max(max_displacement, |delta_p|)\n'
        '\n'
        '    if max_displacement < 0.01:  break         # converged')

    para(doc, 'Force equilibrium point (when attraction = repulsion):')
    code(doc,
        'd_eq = (k_rep / (k_attr * w))^(1/3)\n'
        '     = (100.0 / (0.1 * 0.714))^(1/3)   [kitchen-living pair]\n'
        '     = 1404^(1/3)\n'
        '     = 11.2 m  ->  rooms settle ~11m apart')

    h2(doc, '5.2  Phase B: SLSQP Constraint Optimisation')
    para(doc,
        'scipy.optimize.minimize (method=SLSQP, maxiter=200, ftol=1e-6) refines the '
        'force-directed result. It takes a flat parameter vector x = [x1,y1,w1,l1, x2,y2,w2,l2, ...]')

    code(doc,
        '# layout_agent.py:505\n'
        'def objective(x):\n'
        '    # (1) Adjacency penalty: pull high-weight rooms closer\n'
        '    for i, j where A[i,j] > 0:\n'
        '        score += A[i,j] * max(0, dist(i,j) - min_separation)^2\n'
        '\n'
        '    # (2) Repulsion penalty: push low-weight rooms apart\n'
        '    for i, j where A[i,j] < 0:\n'
        '        score += |A[i,j]| * max(0, 1.5*min_sep - dist(i,j))^2\n'
        '\n'
        '    # (3) Compactness: penalise wasted space in bounding box\n'
        '    bbox_area = (max_x - min_x) * (max_y - min_y)\n'
        '    score += (bbox_area - total_room_area) * 0.01\n'
        '\n'
        '    # (4) Overlap: hard penalty if any two rooms overlap\n'
        '    for each pair (i, j):\n'
        '        overlap = max(0, overlap_x) * max(0, overlap_y)\n'
        '        score += overlap * 1000.0\n'
        '\n'
        '# Constraints (equality, one per room):\n'
        '    width[i] * length[i] == target_area[i]\n'
        '\n'
        '# Bounds per room:\n'
        '    x, y in [0, 100];  width in [min_w, max_w];  length in [min_l, max_l]')

    doc.add_paragraph('')

    # ── 6. STAGE 4: SHAPELY POLYGONS ─────────────────────────────────────────
    h1(doc, '6.  Stage 4 — Shapely Polygon Generation')
    para(doc,
        'After optimisation, each room\'s (x, y, width, length) tuple is converted '
        'to a Shapely Polygon using shapely.geometry.box():')

    code(doc,
        '# layout_agent.py:629\n'
        'from shapely.geometry import box\n'
        '\n'
        'for room_id, pos in optimized_positions.items():\n'
        '    polygon = box(\n'
        '        pos["x"],              # min_x\n'
        '        pos["y"],              # min_y\n'
        '        pos["x"] + pos["width"],   # max_x\n'
        '        pos["y"] + pos["length"]   # max_y\n'
        '    )\n'
        '    room_polygons[room_id] = polygon  # Shapely Polygon object\n'
        '\n'
        '# Shapely Polygon methods used downstream:\n'
        '# polygon.bounds     -> (minx, miny, maxx, maxy)\n'
        '# polygon.centroid   -> Point(cx, cy)  for label placement\n'
        '# polygon.area       -> float\n'
        '# polygon.touches(other)   -> bool  for actual adjacency\n'
        '# polygon.distance(other)  -> float  for near-adjacency (< 0.5m)')

    para(doc, 'Actual adjacency matrix (computed from polygon geometry, not weights):')
    code(doc,
        '# layout_agent.py:834\n'
        'for i, j in room_pairs:\n'
        '    if poly_i.touches(poly_j) or poly_i.intersects(poly_j):\n'
        '        actual_adj[i,j] = 1.0   # physically touching\n'
        '    elif poly_i.distance(poly_j) < 0.5:\n'
        '        actual_adj[i,j] = 0.5   # near-adjacent (within 0.5m)\n'
        '    else:\n'
        '        actual_adj[i,j] = 0.0   # not adjacent')

    doc.add_paragraph('')

    # ── 7. A* CIRCULATION ────────────────────────────────────────────────────
    h1(doc, '7.  Stage 4b — A* Pathfinding for Circulation')
    para(doc,
        'For every room pair where the adjacency weight > 0.3, A* finds the '
        'optimal path through corridor space. These paths become dashed lines '
        'in the SVG floor plan.')

    code(doc,
        '# spatial_algorithms.py:65  AStarPathfinder.find_path()\n'
        '\n'
        'GRID SETUP:\n'
        '  resolution = 0.5 m per cell\n'
        '  grid_bounds = (min_x-2, min_y-2, max_x+2, max_y+2)  [2m margin]\n'
        '  obstacles   = list of (x, y, width, height) per room polygon\n'
        '\n'
        'ALGORITHM:\n'
        '  open_set = MinHeap[(f_score, counter, cell)]\n'
        '  g[start] = 0\n'
        '  f[start] = h(start, goal)  # Manhattan distance heuristic\n'
        '\n'
        '  while open_set:\n'
        '      current = pop_min(open_set)\n'
        '      if current == goal: return reconstruct_path()\n'
        '\n'
        '      for neighbor in 8_directions(current):\n'
        '          if not walkable: skip\n'
        '          tentative_g = g[current] + euclidean(current, neighbor)\n'
        '          if tentative_g < g[neighbor]:\n'
        '              g[neighbor] = tentative_g\n'
        '              f[neighbor] = tentative_g + manhattan(neighbor, goal)\n'
        '              open_set.push((f[neighbor], counter++, neighbor))\n'
        '\n'
        'POST-PROCESSING:\n'
        '  path smoothed via line-of-sight: connect furthest visible waypoints\n'
        '  path_cost = path_length + 0.5 * (num_waypoints - 2)  # turn penalty')

    doc.add_paragraph('')

    # ── 8. SVG FLOOR PLAN ────────────────────────────────────────────────────
    h1(doc, '8.  Stage 5a — SVG Floor Plan  (SVGFloorPlanGenerator)')
    para(doc,
        'SVGFloorPlanGenerator builds an SVG DOM using Python\'s standard '
        'xml.etree.ElementTree, then prettifies it with xml.dom.minidom. '
        'Canvas: 1200 x 800 px. Scale: up to 20 px/m.')

    mk(doc, 3,
        ['SVG Layer', 'Content', 'Implementation'],
        [
            ('1. CSS Styles',       '.room (stroke, fill-opacity), .room-label, .room-area, .circulation, .dimension-text, .title', '_add_styles() -> ET.SubElement(svg, "style")'),
            ('2. Title',            'Floor Plan title centred at top (y=30)', '_add_title() -> ET.SubElement("text")'),
            ('3. Transform',        'scale+translate so all rooms fit within margins; uniform scale = min(scale_x, scale_y, 20px/m)', '_calculate_transform() -> "translate(tx,ty) scale(s,s)"'),
            ('4. Room Polygons',    'Shapely exterior coords -> SVG <polygon points="x1,y1 x2,y2...">; fill from room_colors dict', '_draw_rooms() -> ET.SubElement("polygon")'),
            ('5. Room Labels',      'Room name at polygon.centroid.x, y-0.5; area text at centroid y+0.5', 'centroid from Shapely .centroid property'),
            ('6. Circulation',      'A* path_points -> SVG <polyline> with class="circulation" (blue dashed, opacity=0.6)', '_draw_circulation() -> ET.SubElement("polyline")'),
            ('7. Legend',           'White box at right edge; colour swatch + "name (area m2)" per room', '_add_legend()'),
            ('8. Metrics Overlay',  'White box bottom-left: Total Area, Space Utilization %, Circulation Efficiency %', '_add_metrics_overlay()'),
        ])

    para(doc, 'Room colour palette (SVGFloorPlanGenerator):')
    mk(doc, 3,
        ['Room Type', 'Fill Colour', 'Hex Code'],
        [
            ('bedroom',     'Pale blue',    '#E8F4F8'),
            ('living',      'Pale orange',  '#FFF4E6'),
            ('kitchen',     'Pale green',   '#F0F8E8'),
            ('bathroom',    'Pale purple',  '#F8E8F4'),
            ('dining',      'Pale yellow',  '#FFF8E8'),
            ('circulation', 'Light grey',   '#F0F0F0'),
            ('storage',     'Blue-grey',    '#E8E8F0'),
            ('default',     'White',        '#FFFFFF'),
        ])

    code(doc,
        '# Final SVG output\n'
        'svg_string = SVGFloorPlanGenerator.generate_floor_plan(\n'
        '    layout=layout,\n'
        '    title=f"Floor Plan - {node_id[:8]}",\n'
        '    show_dimensions=True,\n'
        '    show_circulation=True,\n'
        '    show_legend=True\n'
        ')  ->  returns XML string via xml.dom.minidom.parseString().toprettyxml()\n'
        '\n'
        'layout.svg_floor_plan = svg_string   # attached to Layout object')

    doc.add_paragraph('')

    # ── 9. ENHANCED SVG ───────────────────────────────────────────────────────
    h1(doc, '9.  Stage 5b — Enhanced SVG  (EnhancedLayoutVisualizer)')
    para(doc,
        'EnhancedLayoutVisualizer is called after the main layout pipeline completes. '
        'It produces a richer SVG saved to visual_outputs/ and a dual-panel '
        'Matplotlib PNG adjacency graph.')

    h2(doc, '9.1  LayoutRoomAdapter — FBSL to Dict Conversion')
    para(doc,
        'The adapter bridges FBSL Layout objects to the dict format expected '
        'by the enhanced visualizer:')

    code(doc,
        '# enhanced_layout.py:76\n'
        'class LayoutRoomAdapter:\n'
        '    def extract_rooms(self) -> List[Dict]:\n'
        '        for room_id, polygon in self.layout.room_polygons.items():\n'
        '            bounds = polygon.bounds  # (minx, miny, maxx, maxy)\n'
        '            room = self.layout.rooms.get(room_id)\n'
        '            return {\n'
        '                "room_id"     : room_id,\n'
        '                "room_type"   : room.room_type,\n'
        '                "area"        : polygon.area,\n'
        '                "x"           : bounds[0],   # minx\n'
        '                "y"           : bounds[1],   # miny\n'
        '                "width"       : bounds[2] - bounds[0],\n'
        '                "height"      : bounds[3] - bounds[1],\n'
        '                "adjacencies" : room.required_adjacencies,\n'
        '            }')

    h2(doc, '9.2  CompactRoomPlacer — Heuristic Refinement (Fallback)')
    para(doc,
        'If the layout has no room_polygons (e.g. in early GoT expansion), '
        'CompactRoomPlacer arranges rooms by importance score:')

    mk(doc, 2,
        ['Room Type', 'Importance Score'],
        [
            ('living_room / kitchen', '15 / 14'),
            ('hallway / corridor',    '13'),
            ('dining_room',           '12'),
            ('bedroom',               '10'),
            ('bathroom',              '8'),
            ('study / office',        '7'),
            ('utility',               '5'),
            ('storage',               '4'),
            ('garage',                '3'),
        ])

    code(doc,
        '# enhanced_layout.py:136\n'
        'def prioritize(rooms):\n'
        '    def score(room):\n'
        '        base = importance_scores.get(room["room_type"], 5)\n'
        '        adjacency_boost = len(room["adjacencies"]) * 2\n'
        '        area_boost = min(room["area"] / 20.0, 3)\n'
        '        return base + adjacency_boost + area_boost\n'
        '    return sorted(rooms, key=score, reverse=True)\n'
        '\n'
        '# Placement: first room at (0,0), subsequent rooms\n'
        '# try 4 candidate positions around each placed room:\n'
        '#   right, left, below, above (with 0.5m spacing)\n'
        '# pick position with best score = 100 / (dist_to_anchor + 1)')

    h2(doc, '9.3  Enhanced SVG Elements')
    mk(doc, 3,
        ['Element', 'Details', 'Config Parameter'],
        [
            ('Background grid',   '1m x 1m grid lines, colour #ECF0F1, opacity=0.4', 'show_grid=True, grid_spacing=1.0m, grid_color="#ECF0F1"'),
            ('Compass rose',      'Circle + north arrow polygon at top-right; compass_size=60px', 'show_directions=True, compass_size=60.0'),
            ('Room rectangles',   '<rect> per room at scale=30 px/m, wall stroke #2C3E50 width=0.2', 'scale=30.0, wall_color="#2C3E50"'),
            ('Room labels',       'White bold text centred in each room rectangle', 'font_family="Arial", font_size=12px'),
            ('Title bar',         'Project name + node_id[:8] at top-left', 'title_font_size=18px'),
            ('Subtitle',          'Total area (m2) and room count', 'subtitle below title'),
            ('Output file',       'visual_outputs/{project}_{node_id[:8]}_layout_{timestamp}.svg', 'VisualConfig.output_root'),
        ])

    para(doc, 'Enhanced room colour palette (richer, dark colours):')
    mk(doc, 3,
        ['Room Type', 'Colour Name', 'Hex'],
        [
            ('living_room',  'Purple',       '#9B59B6'),
            ('dining_room',  'Orange',       '#F39C12'),
            ('kitchen',      'Green',        '#27AE60'),
            ('bedroom',      'Blue',         '#3498DB'),
            ('bathroom',     'Burnt orange', '#E67E22'),
            ('study/office', 'Red',          '#E74C3C'),
            ('corridor/hall','Silver',       '#BDC3C7'),
            ('utility',      'Teal',         '#16A085'),
            ('storage',      'Dark purple',  '#8E44AD'),
            ('garage',       'Grey',         '#95A5A6'),
            ('default',      'Dark slate',   '#34495E'),
        ])

    doc.add_paragraph('')

    # ── 10. ADJACENCY GRAPH PNG ────────────────────────────────────────────────
    h1(doc, '10.  Adjacency Graph PNG  (Matplotlib + NetworkX)')
    para(doc,
        'The dual-panel adjacency graph is saved as a 200 DPI PNG. '
        'It combines a geometric floor plot (left) with a spring-layout graph (right) '
        'in a single 18x9 inch Matplotlib figure.')

    h2(doc, '10.1  NetworkX Graph Construction')
    code(doc,
        '# enhanced_layout.py:391\n'
        'graph = nx.Graph()\n'
        '\n'
        'for room in rooms:\n'
        '    graph.add_node(room["room_id"],\n'
        '        room_type=room["room_type"],\n'
        '        area=room["area"],\n'
        '        x=room["x"], y=room["y"],\n'
        '        width=room["width"], height=room["height"])\n'
        '\n'
        '# Edge detection: geometry-based (not weight-based)\n'
        'for each pair (room_a, room_b):\n'
        '    if rooms_adjacent(room_a, room_b):           # shared wall check\n'
        '        edge_type = _edge_type(room_a, room_b)   # critical/spatial\n'
        '        graph.add_edge(room_a_id, room_b_id, edge_type=edge_type)\n'
        '\n'
        '# Connectivity guarantee: if graph is disconnected,\n'
        '# add "bridge" edges between components\n'
        'if not nx.is_connected(graph):\n'
        '    for each disconnected component:\n'
        '        add_edge(comp_node, main_component_node, edge_type="bridge")')

    h2(doc, '10.2  Edge Classification')
    mk(doc, 4,
        ['Edge Type', 'Condition', 'Colour', 'Line Style'],
        [
            ('critical',   'kitchen<->dining/living, bedroom<->bathroom/corridor', '#27AE60 (green)', 'solid, width=3.5'),
            ('spatial',    'any other geometry-detected adjacency',                '#3498DB (blue)',  'solid, width=2.5'),
            ('proximity',  'fallback: nearest N-1 pairs by centroid distance',     '#95A5A6 (grey)',  'dotted, width=1.5'),
            ('bridge',     'added to reconnect disconnected components',            '#E74C3C (red)',   'dashed, width=2.0'),
        ])

    code(doc,
        '# _edge_type()  enhanced_layout.py:430\n'
        'critical_pairs = {\n'
        '    "kitchen":  {"dining_room", "living_room"},\n'
        '    "bedroom":  {"bathroom", "corridor", "hallway"},\n'
        '}\n'
        'if room_b_type in critical_pairs.get(room_a_type, set()): return "critical"\n'
        'if room_a_type in critical_pairs.get(room_b_type, set()): return "critical"\n'
        'return "spatial"')

    h2(doc, '10.3  Geometry-Based Adjacency Detection')
    code(doc,
        '# _rooms_adjacent()  enhanced_layout.py:465\n'
        'tolerance = 0.2  # meters\n'
        '\n'
        'horizontal = |right1 - left2| <= tol  OR  |right2 - left1| <= tol\n'
        'vertical   = |bottom1 - top2| <= tol  OR  |bottom2 - top1| <= tol\n'
        '\n'
        'if horizontal:\n'
        '    overlap = min(bottom1, bottom2) - max(top1, top2)  # shared wall height\n'
        '    return overlap >= 0.3 * min(height1, height2)      # 30% overlap minimum\n'
        '\n'
        'if vertical:\n'
        '    overlap = min(right1, right2) - max(left1, left2)  # shared wall width\n'
        '    return overlap >= 0.3 * min(width1, width2)\n'
        '\n'
        '# If no physical adjacency found: fallback to proximity edges\n'
        '# (sorted by centroid distance, add nearest len(rooms)-1 pairs)')

    h2(doc, '10.4  Dual-Panel Matplotlib Figure')
    code(doc,
        '# _generate_adjacency_graph()  enhanced_layout.py:405\n'
        'fig, (ax_layout, ax_graph) = plt.subplots(1, 2, figsize=(18, 9))\n'
        'fig.suptitle(f"{project_name} - Connectivity ({node_id[:8]})",\n'
        '             fontsize=16, fontweight="bold")\n'
        '\n'
        '# LEFT PANEL: _plot_spatial(ax_layout, graph)\n'
        '#   Draw room rectangles as matplotlib.patches.Rectangle\n'
        '#   Room fill = config.room_colors[room_type]\n'
        '#   White text label centred in each room\n'
        '#   Draw lines between adjacent rooms (centroid to centroid)\n'
        '#   ax.invert_yaxis()  # match screen coordinates\n'
        '#   ax.set_aspect("equal")\n'
        '\n'
        '# RIGHT PANEL: _plot_graph(ax_graph, graph)\n'
        '#   pos = nx.spring_layout(graph, seed=42)\n'
        '#   nx.draw_networkx_nodes(pos, node_color=room_colors, node_size=900)\n'
        '#   Draw edges by type with different colour/width/style\n'
        '#   nx.draw_networkx_labels(pos, font_size=8, font_color="white")\n'
        '#   legend: critical/spatial/proximity/bridge\n'
        '\n'
        'fig.tight_layout()\n'
        'fig.savefig(graph_path, dpi=200)  # 200 DPI -> high resolution PNG\n'
        'plt.close(fig)')

    doc.add_paragraph('')

    # ── 11. AdjacencyGraphVisualizer ──────────────────────────────────────────
    h1(doc, '11.  Legacy Adjacency Graph SVG  (AdjacencyGraphVisualizer)')
    para(doc,
        'A second adjacency graph is generated as an SVG (not PNG) by '
        'AdjacencyGraphVisualizer in utils/visualization.py. '
        'This is stored on layout.adjacency_svg and uses NetworkX spring_layout '
        'to compute positions, then renders nodes and edges as SVG circles and lines.')

    code(doc,
        '# layout_agent.py:228\n'
        'room_ids_list = list(node.layout.rooms.keys())\n'
        '\n'
        'adjacency_graph = _build_networkx_adjacency_graph(\n'
        '    room_polygons,      # {room_id: Shapely Polygon}\n'
        '    adjacency_matrix,   # 12x12 NumPy ndarray (weight-based)\n'
        '    room_ids_list       # ordered list\n'
        ')\n'
        '\n'
        '# Nodes: each room_id with area and centroid attributes\n'
        '# Edges: all pairs where adjacency_matrix[i,j] > 0.1\n'
        '#        edge.weight = adjacency_matrix[i,j]\n'
        '\n'
        'adjacency_svg = graph_visualizer.generate_adjacency_svg(\n'
        '    adjacency_graph=adjacency_graph,\n'
        '    title=f"Adjacency Graph - {node_id[:8]}",\n'
        '    show_weights=True,\n'
        '    layout_type="spring"    # nx.spring_layout()\n'
        ')\n'
        'layout.adjacency_svg = adjacency_svg')

    doc.add_paragraph('')

    # ── 12. OUTPUT FILES ──────────────────────────────────────────────────────
    h1(doc, '12.  Output Files Summary')

    mk(doc, 5,
        ['Output', 'Format', 'Location', 'Size', 'Generated By'],
        [
            ('Floor plan SVG (standard)',  'SVG string',  'layout.svg_floor_plan (in-memory)', '~50 KB',  'SVGFloorPlanGenerator'),
            ('Enhanced floor plan SVG',    'SVG file',    'visual_outputs/{project}_{node}_{ts}.svg', '~30 KB', 'EnhancedLayoutVisualizer'),
            ('Adjacency graph PNG',        'PNG raster',  'visual_outputs/{project}_{node}_adjacency_{ts}.png', '~200 KB @ 200 DPI', 'EnhancedLayoutVisualizer + Matplotlib'),
            ('Adjacency graph SVG',        'SVG string',  'layout.adjacency_svg (in-memory)', '~20 KB',  'AdjacencyGraphVisualizer (NetworkX)'),
            ('Prototype JSON',             'JSON',        'pipeline_outputs/prototype_{n}.json', '~1 KB',  'FinalOutputAgent'),
            ('Pareto analysis',            'TXT',         'pipeline_outputs/pareto_analysis.txt', '~2 KB', 'AggregationAgent'),
        ])

    doc.add_paragraph('')

    # ── 13. FULL DATA FLOW DIAGRAM ────────────────────────────────────────────
    h1(doc, '13.  Complete Data Flow Summary')

    mk(doc, 3,
        ['Step', 'Data Object', 'Libraries Involved'],
        [
            ('NL Text -> FBSL Node',          'FBSLLayoutNode (functions, behaviors, structures, rooms)', 'Ollama LLM, pydantic'),
            ('FBSL Node -> Room Specs',        'Dict[room_id, {area, width, length}]', 'NumPy (sqrt)'),
            ('Room Specs -> Adj Matrix',       'NumPy ndarray [12x12]', 'NumPy, WeightedAdjacencyCalculator'),
            ('Adj Matrix -> Initial Positions','Dict[room_id, {x, y}]', 'NumPy, random, ForceDirectedLayout'),
            ('Positions -> Refined Positions', 'Dict[room_id, {x, y, width, length}]', 'scipy.optimize.minimize (SLSQP)'),
            ('Positions -> Shapely Polygons',  'Dict[room_id, Shapely Polygon]', 'Shapely (box())'),
            ('Polygons -> Circulation Paths',  'List[CirculationPath]', 'heapq, AStarPathfinder (0.5m grid)'),
            ('Polygons + Paths -> Layout',     'Layout object (metrics computed)', 'Shapely, NumPy, NetworkX'),
            ('Layout -> SVG Floor Plan',       'SVG XML string',  'xml.etree.ElementTree, xml.dom.minidom'),
            ('Layout -> Enhanced SVG',         'SVG file on disk', 'xml.etree.ElementTree, Shapely'),
            ('Layout -> Adj Graph PNG',        'PNG file on disk', 'NetworkX (spring_layout), Matplotlib (200 DPI)'),
            ('Layout -> Score',                'Dict[criterion, float]', 'NumPy (geometric mean, harmonic mean)'),
        ])

    doc.add_paragraph('')

    # ── SAVE ──────────────────────────────────────────────────────────────────
    candidates = [
        Path(r'c:\Users\sanvi\OneDrive\Desktop\layout\docs\FBSL_KAGS_Layout_Visualization_Pipeline.docx'),
        Path(r'c:\Users\sanvi\Documents\FBSL_KAGS_Layout_Visualization_Pipeline.docx'),
        Path(r'c:\Users\sanvi\FBSL_Layout_Viz.docx'),
    ]
    for out in candidates:
        try:
            out.parent.mkdir(parents=True, exist_ok=True)
            doc.save(str(out))
            print(f"Saved: {out}")
            print(f"Size : {out.stat().st_size / 1024:.1f} KB")
            break
        except Exception as e:
            print(f"Could not save to {out}: {e}", file=sys.stderr)


if __name__ == '__main__':
    try:
        build()
    except Exception as e:
        print(f"BUILD FAILED: {e}", file=sys.stderr)
        traceback.print_exc()
        sys.exit(1)
