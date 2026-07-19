from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_TABLE_ALIGNMENT

doc = Document()

# Set up styles
style = doc.styles['Title']
style.font.size = Pt(24)
style.font.bold = True
style.font.color.rgb = RGBColor(0, 51, 102)

style = doc.styles['Heading 1']
style.font.size = Pt(16)
style.font.bold = True
style.font.color.rgb = RGBColor(0, 51, 102)

style = doc.styles['Heading 2']
style.font.size = Pt(14)
style.font.bold = True
style.font.color.rgb = RGBColor(0, 102, 153)

style = doc.styles['Heading 3']
style.font.size = Pt(12)
style.font.bold = True
style.font.color.rgb = RGBColor(51, 51, 51)

# Title
title = doc.add_heading('FBSL-KAGS: Complete System Documentation', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph()

# ============ SECTION 1: AGENT SYSTEM ============
doc.add_heading('1. AGENT SYSTEM (9 Agents)', level=1)

# Agent 1.1
doc.add_heading('1.1 Encoder Agent', level=2)
doc.add_paragraph('Role: The Encoder Agent serves as the entry point to the system. It transforms unstructured natural language requirements into a structured FBSL (Function-Behavior-Structure-Layout) representation that the rest of the pipeline can process.')
doc.add_paragraph('Input: Raw user text describing their design requirements (e.g., "Design a 3-bedroom house with open kitchen connected to living room, good natural lighting, and quiet bedrooms")')
doc.add_paragraph('Output: An initial FBSL problem node containing:')
p = doc.add_paragraph()
p.add_run('• Functions (F): ').bold = True
p.add_run('Extracted room types with priorities and activities')
p = doc.add_paragraph()
p.add_run('• Expected Behaviors (Bₑ): ').bold = True
p.add_run('Derived performance requirements (thermal, acoustic, lighting)')
p = doc.add_paragraph()
p.add_run('• Initial Structures (S): ').bold = True
p.add_run('Placeholder structural elements')
p = doc.add_paragraph()
p.add_run('• Initial Layout (L): ').bold = True
p.add_run('Basic room list without positioning')

doc.add_paragraph('LLM Integration: Uses Ollama (gemma3 or llama3.1) to parse natural language. The LLM receives the user text and returns structured JSON containing room types, counts, spatial constraints, adjacency requirements, and behavioral expectations.')

doc.add_paragraph('Key Operations:')
doc.add_paragraph('• Extracts room types from keywords ("bedroom", "kitchen", "bathroom")')
doc.add_paragraph('• Identifies numerical constraints ("3 bedrooms", "25 sqm living room")')
doc.add_paragraph('• Detects adjacency requirements ("kitchen connected to dining")')
doc.add_paragraph('• Derives behavioral expectations from qualitative terms ("quiet" → acoustic requirements)')
doc.add_paragraph('• Assigns function priorities based on emphasis in user description')

# Agent 1.2
doc.add_heading('1.2 Generalizer Agent', level=2)
doc.add_paragraph('Role: The Generalizer Agent takes the encoded problem and decomposes it into multiple design alternatives. This creates diversity in the design space, ensuring the system explores different approaches rather than converging prematurely on a single solution.')
doc.add_paragraph('Input: The encoded FBSL problem node from the Encoder Agent')
doc.add_paragraph('Output: 4 alternative problem variants, each representing a different design philosophy')

doc.add_paragraph('Four Decomposition Strategies:', style='Heading 3')
doc.add_paragraph('1. Functional Zone Decomposition: Groups rooms by zone type (private/social/service). Each variant emphasizes different zone relationships.')
doc.add_paragraph('2. Topology Variants: Different spatial organization patterns - Linear, Clustered, Grid, or Radial arrangements.')
doc.add_paragraph('3. Priority-Based Variants: Different function emphasis (maximize living space, optimize bedroom privacy, enhance kitchen functionality, or balance all).')
doc.add_paragraph('4. Structural Alternatives: Different construction approaches (heavy construction, light construction, or mixed systems).')

# Agent 1.3
doc.add_heading('1.3 Research Agent (External RAG)', level=2)
doc.add_paragraph('Role: The Research Agent interfaces with the external knowledge base to retrieve relevant precedent designs. It enhances the current design context with proven patterns from the Finnish floor plan database, enabling knowledge-augmented generation.')
doc.add_paragraph('Input: Current FBSL context (functions, behaviors, spatial requirements)')
doc.add_paragraph('Output: Top-k similar precedent designs with metadata (typically k = 3-5 per function)')

doc.add_paragraph('FAISS Vector Store Integration:', style='Heading 3')
doc.add_paragraph('• The Research Agent queries the FAISS index containing embeddings of Finnish floor plans')
doc.add_paragraph('• Uses IndexFlatIP (Inner Product) for exact cosine similarity search')
doc.add_paragraph('• Query process: Convert context to embedding → Normalize vector → Search index → Map indices to metadata')

p = doc.add_paragraph()
p.add_run('Similarity Formula: ').bold = True
p.add_run('similarity(q, pᵢ) = cos(E(q), E(pᵢ)) = (q · pᵢ) / (||q|| × ||pᵢ||)')
doc.add_paragraph('This cosine similarity measures the angle between the query embedding and each precedent embedding. Values close to 1.0 indicate high semantic similarity.')

doc.add_paragraph('Retrieval Thresholds by Priority:')
doc.add_paragraph('• High priority functions: similarity > 0.7 required')
doc.add_paragraph('• Medium priority functions: similarity > 0.5 required')
doc.add_paragraph('• Low priority functions: similarity > 0.3 required')

# Agent 1.4
doc.add_heading('1.4 Scoring Agent', level=2)
doc.add_paragraph('Role: The Scoring Agent evaluates the quality of each design alternative across five dimensions using Multi-Criteria Decision Analysis (MCDA). It provides a unified quality metric that enables comparison and selection of designs.')
doc.add_paragraph('Input: Complete FBSL node with all components populated')
doc.add_paragraph('Output: Five individual scores plus a composite score ∈ [0, 1]')

doc.add_paragraph('Five Evaluation Dimensions:', style='Heading 3')

p = doc.add_paragraph()
p.add_run('1. Functional Adequacy (S_f) - Weight: 0.30').bold = True
doc.add_paragraph('• Measures how well the design satisfies stated functions')
doc.add_paragraph('• Formula: S_f = Σᵢ (wᵢ × Coverage(fᵢ)) / Σᵢ wᵢ')
doc.add_paragraph('• Coverage(fᵢ) = Σⱼ min(1, Bₛⱼ/Bₑⱼ) / m where m = number of related behaviors')
doc.add_paragraph('• Higher function priority (wᵢ) means that function contributes more to the score')

p = doc.add_paragraph()
p.add_run('2. Behavioral Performance (S_b) - Weight: 0.30').bold = True
doc.add_paragraph('• Measures how closely actual behaviors match expected behaviors')
doc.add_paragraph('• Uses geometric mean to ensure ALL behaviors must perform reasonably')
doc.add_paragraph('• Formula: S_b = [Πᵢ min(1, Bₛᵢ/Bₑᵢ)]^(1/m)')
doc.add_paragraph('• Why geometric mean? One failing behavior (e.g., 0.1) pulls down the entire score, unlike arithmetic mean which allows good behaviors to mask failures')

p = doc.add_paragraph()
p.add_run('3. Structural Feasibility (S_s) - Weight: 0.20').bold = True
doc.add_paragraph('• Checks if proposed structures are physically valid and compatible')
doc.add_paragraph('• Formula: S_s = Σⱼ (Feasibility(sⱼ) × Compatibility(sⱼ, others)) / |S|')
doc.add_paragraph('• Feasibility = material validity × dimensional constraints')
doc.add_paragraph('• Compatibility = 1 - (conflicts / total_structures)')

p = doc.add_paragraph()
p.add_run('4. Layout Efficiency (S_l) - Weight: 0.15').bold = True
doc.add_paragraph('• Evaluates spatial arrangement quality')
doc.add_paragraph('• Formula: S_l = 0.4×Compactness + 0.3×Circulation + 0.3×Adjacency')
doc.add_paragraph('• Compactness = Total_Room_Area / Bounding_Box_Area')
doc.add_paragraph('• Circulation = Direct_Distance / Actual_Path_Length')
doc.add_paragraph('• Adjacency = Satisfied_Requirements / Total_Requirements')

p = doc.add_paragraph()
p.add_run('5. Sustainability (S_sust) - Weight: 0.05').bold = True
doc.add_paragraph('• Considers environmental impact')
doc.add_paragraph('• Formula: S_sust = Σᵢ (wᵢ × sustainability_metricᵢ)')
doc.add_paragraph('• Includes energy efficiency, material sustainability, environmental impact')

p = doc.add_paragraph()
p.add_run('Composite Score: ').bold = True
p.add_run('S_composite = 0.30×S_f + 0.30×S_b + 0.20×S_s + 0.15×S_l + 0.05×S_sust')

# Agent 1.5
doc.add_heading('1.5 Layout Generation Agent', level=2)
doc.add_paragraph('Role: The Layout Generation Agent creates actual spatial arrangements by positioning rooms and generating circulation paths. It uses physics-based optimization to find layouts that satisfy adjacency requirements while avoiding overlaps.')
doc.add_paragraph('Input: Room list with dimensions and adjacency matrix from FBSL')
doc.add_paragraph('Output: Complete Layout containing positions, dimensions, circulation paths, and layout metrics')

doc.add_paragraph('Weighted Adjacency Matrix:', style='Heading 3')
doc.add_paragraph('• Defines spatial relationship preferences between room pairs')
doc.add_paragraph('• Formula: w(i,j) = 0.4×Functional_Dependency + 0.35×Traffic_Flow + 0.25×Privacy')
doc.add_paragraph('• Values ∈ [-1, 1]: Positive = rooms should be adjacent, Negative = rooms should be separated')

doc.add_paragraph('Force-Directed Layout Algorithm:', style='Heading 3')
doc.add_paragraph('• Treats rooms as particles in a physics simulation')
doc.add_paragraph('• Attractive forces pull adjacent rooms together (like springs)')
doc.add_paragraph('• Repulsive forces prevent room overlaps (like electric charges)')
p = doc.add_paragraph()
p.add_run('Force formula: ').bold = True
p.add_run('Force(rᵢ, rⱼ) = k_attraction × A[i,j] × d(rᵢ, rⱼ) - k_repulsion / d²(rᵢ, rⱼ)')
doc.add_paragraph('  • k_attraction = 0.1 (spring constant)')
doc.add_paragraph('  • k_repulsion = 100.0 (charge strength)')
doc.add_paragraph('  • A[i,j] = adjacency weight from matrix')
doc.add_paragraph('  • d = Euclidean distance between room centers')
p = doc.add_paragraph()
p.add_run('Position update: ').bold = True
p.add_run('pᵢ(t+1) = pᵢ(t) + η × ΣForces where η = 0.01 (learning rate)')
doc.add_paragraph('• Convergence: Stop when max displacement < 0.01m or iterations ≥ 200')

doc.add_paragraph('A* Pathfinding for Circulation:', style='Heading 3')
doc.add_paragraph('• Generates optimal walking paths between rooms')
p = doc.add_paragraph()
p.add_run('Cost function: ').bold = True
p.add_run('f(n) = g(n) + h(n)')
doc.add_paragraph('  • g(n) = actual distance traveled from start (Euclidean)')
doc.add_paragraph('  • h(n) = estimated distance to goal (Manhattan heuristic)')
doc.add_paragraph('• Process: Create 0.5m grid → Mark rooms as obstacles → Find shortest path → Smooth path')

# Agent 1.6
doc.add_heading('1.6 Refinement Agent', level=2)
doc.add_paragraph('Role: The Refinement Agent iteratively improves designs by comparing actual behaviors (Bₛ) against expected behaviors (Bₑ) and applying appropriate corrections. This is based on Gero\'s FBS reformulation framework with three types of modifications.')
doc.add_paragraph('Input: FBSL node where actual behaviors don\'t match expected behaviors')
doc.add_paragraph('Output: Refined FBSL node with improved behavior satisfaction')

p = doc.add_paragraph()
p.add_run('Deviation Calculation: ').bold = True
p.add_run('avg_deviation = mean(|Bₛᵢ - Bₑᵢ| / Bₑᵢ) for all behaviors')
doc.add_paragraph('This measures the average percentage gap between what we wanted and what we achieved.')

doc.add_heading('Three Reformulation Types', level=3)

p = doc.add_paragraph()
p.add_run('Type 1: Structure Modification (when avg_deviation < 0.3)').bold = True
doc.add_paragraph('• Situation: Small gap - we\'re close to meeting targets')
doc.add_paragraph('• Action: Add or modify structures to close the behavior gap')
doc.add_paragraph('• Formula: S\' = S + ΔS where ΔS minimizes |Bₛ - Bₑ|')
doc.add_paragraph('• Candidate Structures by Behavior Category:')
doc.add_paragraph('  - THERMAL not satisfied → Add insulation layer, upgrade wall material')
doc.add_paragraph('  - ACOUSTIC not satisfied → Add sound partition (STC 50+)')
doc.add_paragraph('  - LIGHTING not satisfied → Add window opening (1.5m × 2.0m)')
doc.add_paragraph('  - VENTILATION not satisfied → Add MEP duct system')

p = doc.add_paragraph()
p.add_run('Type 2: Behavior Relaxation (when 0.3 ≤ avg_deviation < 0.6)').bold = True
doc.add_paragraph('• Situation: Moderate gap - meeting original targets may be too difficult')
doc.add_paragraph('• Action: Expand acceptable performance ranges')
doc.add_paragraph('• Formula: tolerance\' = tolerance × 1.2 (increase by 20%)')
doc.add_paragraph('• Example: Temperature target = 21°C, tolerance 10% → Range [18.9°C, 23.1°C]')
doc.add_paragraph('  After relaxation: Tolerance = 12% → Range [18.5°C, 23.5°C]')

p = doc.add_paragraph()
p.add_run('Type 3: Function Redefinition (when avg_deviation ≥ 0.6)').bold = True
doc.add_paragraph('• Situation: Large gap - the problem is fundamentally too constrained')
doc.add_paragraph('• Action: Reduce priority of problematic functions')
doc.add_paragraph('• Formula: priority\' = priority × 0.8')
doc.add_paragraph('• Process: Identify low-priority functions causing largest deviations → Reduce priority by 20% → System focuses on achievable high-priority functions')

doc.add_heading('Refinement Loop Workflow', level=3)
doc.add_paragraph('1. Calculate Bₛ from current structures S (using physics-based calculation)')
doc.add_paragraph('2. Identify unsatisfied behaviors (compare each Bₛᵢ to Bₑᵢ)')
doc.add_paragraph('3. Calculate composite_score using all 5 dimensions')
doc.add_paragraph('4. Check convergence: |score - prev_score| < 0.01 → If converged: EXIT loop')
doc.add_paragraph('5. Calculate avg_deviation across unsatisfied behaviors')
doc.add_paragraph('6. Apply appropriate reformulation (Type 1, 2, or 3 based on deviation)')
doc.add_paragraph('7. Update FBSL node with changes')
doc.add_paragraph('8. Loop back to Step 1 with modified node')

doc.add_heading('How Behavior Improvement Works', level=3)
doc.add_paragraph('Step-by-step example showing temperature improvement:')
doc.add_paragraph('1. Initial State: Expected temp = 21°C, Structure has U-value = 0.5 W/m²K, Achieves 18°C (deviation = 14%)')
doc.add_paragraph('2. Iteration 1 (Type 1): Add insulation → New U-value = 0.25 W/m²K → Achieves 20°C (deviation = 5%)')
doc.add_paragraph('3. Iteration 2 (Type 1): Add double glazing → Achieves 20.5°C (deviation = 2.4%)')
doc.add_paragraph('4. Iteration 3: Minor adjustment → Achieves 20.8°C (within tolerance)')
doc.add_paragraph('5. Convergence: Score improvement < 0.01 → CONVERGED')
doc.add_paragraph('The improvement happens because each structure modification directly affects the physics: more insulation → higher R-value → better temperature maintenance.')

# Agent 1.7
doc.add_heading('1.7 Pruning Agent', level=2)
doc.add_paragraph('Role: The Pruning Agent removes low-quality design nodes from the Graph of Thoughts exploration. This focuses computational resources on promising design directions and prevents the search from wasting time on inferior alternatives.')
doc.add_paragraph('Input: All scored nodes from GoT exploration')
doc.add_paragraph('Output: Filtered set containing only high-potential candidates')

p = doc.add_paragraph()
p.add_run('Pruning Threshold: ').bold = True
p.add_run('pruning_threshold = max_score × 0.70')
doc.add_paragraph('Any node with composite_score below 70% of the best score is pruned.')

doc.add_paragraph('Pareto Preservation Rule:')
doc.add_paragraph('• After pruning, check the Pareto frontier size')
doc.add_paragraph('• If |Pareto_set| < min_pareto_solutions (default: 3): Restore highest-scoring dominated solutions')
doc.add_paragraph('• This ensures diversity is maintained even after aggressive pruning')

# Agent 1.8
doc.add_heading('1.8 Aggregation Agent', level=2)
doc.add_paragraph('Role: The Aggregation Agent combines multiple high-scoring designs into an improved composite design. When several alternatives each have unique strengths, aggregation can create a design that captures the best features of each.')
doc.add_paragraph('Input: High-scoring nodes (score ≥ max × 0.9)')
doc.add_paragraph('Output: Merged composite design OR individual nodes if incompatible')

p = doc.add_paragraph()
p.add_run('Trigger Condition: ').bold = True
p.add_run('IF (high_scoring_count ≥ 2) AND (avg_compatibility > 0.7): Perform aggregation')

p = doc.add_paragraph()
p.add_run('Compatibility Formula: ').bold = True
p.add_run('Compatibility(Nᵢ, Nⱼ) = 1 - (total_conflicts / total_elements)')

doc.add_paragraph('Conflict Detection:')
doc.add_paragraph('• Function conflicts: f₁.conflicts_with contains f₂')
doc.add_paragraph('• Behavior conflicts: |target₁ - target₂| / max(targets) > 0.5')
doc.add_paragraph('• Structure conflicts: Incompatible materials (wood + heavy concrete)')
doc.add_paragraph('• Layout conflicts: Area deviation > 20%, room count difference > 2')

p = doc.add_paragraph()
p.add_run('Aggregation Formula: ').bold = True
p.add_run('Aggregate(N₁, N₂, ..., Nₖ) = argmax[Σᵢ λᵢ × Compatibility × Quality(Nᵢ)]')

# Agent 1.9
doc.add_heading('1.9 Pipeline Orchestrator', level=2)
doc.add_paragraph('Role: The Pipeline Orchestrator is the central coordinator that manages the entire design generation process. It controls the Graph of Thoughts exploration, invokes agents in the correct sequence, and determines when to stop exploration.')
doc.add_paragraph('Input: Original user requirements')
doc.add_paragraph('Output: Top 3-5 design prototypes ranked by quality')

p = doc.add_paragraph()
p.add_run('Complexity Calculation: ').bold = True
p.add_run('C_overall = 0.4 × C_req + 0.6 × C_fbsl')

doc.add_paragraph('Complexity-Based Scaling:')
doc.add_paragraph('• Low (< 0.3): Scale = 0.7 → Shallow search, fewer nodes')
doc.add_paragraph('• Medium (0.3-0.6): Scale = 1.0 → Standard parameters')
doc.add_paragraph('• High (0.6-0.8): Scale = 1.3 → Deeper search, more alternatives')
doc.add_paragraph('• Very High (≥ 0.8): Scale = 1.5 → Maximum exploration')

p = doc.add_paragraph()
p.add_run('GoT Stopping Criteria: ').bold = True
doc.add_paragraph('Stop if: (improvement < 0.001) AND (score > 0.7) AND (stagnation ≥ 2)')
doc.add_paragraph('OR high_scoring_count ≥ 3')
doc.add_paragraph('OR Pareto_frontier_stable')
doc.add_paragraph('OR depth/node_limit_reached')

# ============ SECTION 2: PIPELINE WORKFLOW ============
doc.add_heading('2. OVERALL PIPELINE WORKFLOW', level=1)

doc.add_heading('Phase 0: Input Reception', level=2)
doc.add_paragraph('User provides natural language requirements describing their design needs.')

doc.add_heading('Phase 1: Encoding & Knowledge Retrieval', level=2)
doc.add_paragraph('Step 1: Requirement Parsing (Encoder Agent) - LLM receives user text, extracts rooms, constraints, behaviors, priorities, creates initial FBSL problem node.')
doc.add_paragraph('Step 2: Precedent Retrieval (Research Agent) - Encodes problem context, queries FAISS, retrieves similar Finnish floor plans, extracts successful patterns.')
doc.add_paragraph('Step 3: Knowledge Enhancement - Augments problem node with precedent insights, validates room sizes, adds precedent-derived behavior targets.')

doc.add_heading('Phase 2: Design Space Generation', level=2)
doc.add_paragraph('Step 4: Alternative Generation (Generalizer Agent) - Creates 4 design variants using different strategies.')
doc.add_paragraph('Step 5: Complexity Analysis (Orchestrator) - Calculates C_overall, determines adaptive parameters.')
doc.add_paragraph('Step 6: Graph of Thoughts Expansion (Orchestrator) - Builds design graph using 4 expansion strategies.')

doc.add_heading('Phase 3: Evaluation & Selection', level=2)
doc.add_paragraph('Step 7: Comprehensive Scoring (Scoring Agent) - Calculates all 5 scores for every node.')
doc.add_paragraph('Step 8: Low-Quality Pruning (Pruning Agent) - Removes nodes with score < 70% of maximum, preserves Pareto frontier.')
doc.add_paragraph('Step 9: High-Quality Aggregation (Aggregation Agent) - Identifies high-scorers, checks compatibility, merges if conditions met.')

doc.add_heading('Phase 4: Refinement & Optimization', level=2)
doc.add_paragraph('Step 10: Iterative Refinement (Refinement Agent) - For each candidate: calculate Bₛ, compare to Bₑ, apply reformulation, repeat until convergence.')
doc.add_paragraph('Step 11: Spatial Layout Generation (Layout Agent) - Build adjacency matrix, run force-directed optimization, generate circulation paths.')

doc.add_heading('Phase 5: Final Output', level=2)
doc.add_paragraph('Step 12: Final Ranking (Scoring Agent) - Re-score all refined designs, apply Pareto analysis.')
doc.add_paragraph('Step 13: Prototype Selection (Orchestrator) - Select top 3-5 designs by composite score.')
doc.add_paragraph('Step 14: Result Packaging - Each prototype includes complete FBSL, all scores, visualization, refinement history, trade-off analysis.')

# ============ SECTION 3: EXTERNAL RAG ============
doc.add_heading('3. EXTERNAL RAG SYSTEM', level=1)

doc.add_heading('3.1 FAISS Vector Store', level=2)
doc.add_paragraph('Purpose: Enables fast similarity search over thousands of Finnish floor plan embeddings.')
doc.add_paragraph('Index Configuration: IndexFlatIP (Inner Product) - After L2 normalization, inner product equals cosine similarity.')
doc.add_paragraph('Query Process:')
doc.add_paragraph('1. Receive query from Research Agent')
doc.add_paragraph('2. Convert query to embedding vector using same encoder')
doc.add_paragraph('3. Normalize: faiss.normalize_L2(query_vector)')
doc.add_paragraph('4. Search: distances, indices = index.search(query_vector, top_k)')
doc.add_paragraph('5. Map indices to metadata dictionary')

doc.add_heading('3.2 Five Embedding Types', level=2)

p = doc.add_paragraph()
p.add_run('1. Text Embeddings (Weight: 30%)').bold = True
doc.add_paragraph('• Source: Room labels, Finnish architectural terms')
doc.add_paragraph('• Model: Sentence Transformer (all-MiniLM-L6-v2)')
doc.add_paragraph('• Purpose: Semantic understanding of room descriptions')

p = doc.add_paragraph()
p.add_run('2. Architectural Embeddings (Weight: 40%)').bold = True
doc.add_paragraph('• Source: Domain-specific features')
doc.add_paragraph('• Components: Function hierarchy, privacy zones, natural light access')
doc.add_paragraph('• Purpose: Capture architectural domain knowledge')

p = doc.add_paragraph()
p.add_run('3. Spatial Embeddings (Weight: 20%)').bold = True
doc.add_paragraph('• Source: Geometric coordinates and relationships')
doc.add_paragraph('• Components: Normalized positions, adjacency graphs, proximity to edges')
doc.add_paragraph('• Purpose: Encode spatial arrangement patterns')

p = doc.add_paragraph()
p.add_run('4. Visual Embeddings (Weight: 10%)').bold = True
doc.add_paragraph('• Source: Floor plan images')
doc.add_paragraph('• Model: CLIP ViT-B/32 (vision encoder)')
doc.add_paragraph('• Purpose: Visual similarity matching')

p = doc.add_paragraph()
p.add_run('5. Composite Embeddings (Primary)').bold = True
doc.add_paragraph('• Fusion: E_composite = 0.3×E_text + 0.4×E_arch + 0.2×E_spatial + 0.1×E_visual')
doc.add_paragraph('• Purpose: Unified multi-modal representation for primary RAG retrieval')

# ============ SECTION 4: LLM ROLE ============
doc.add_heading('4. LLM ROLE (Ollama)', level=1)

doc.add_heading('4.1 Models Used', level=2)
doc.add_paragraph('Primary: gemma3 or llama3.1 (running locally via Ollama)')
doc.add_paragraph('Purpose: Natural language understanding, structured extraction, reasoning')

doc.add_heading('4.2 LLM Integration Points', level=2)

p = doc.add_paragraph()
p.add_run('In Encoder Agent: ').bold = True
p.add_run('Parse unstructured requirements into structured FBSL. Input: User\'s natural language. Output: JSON with rooms, constraints, behaviors.')

p = doc.add_paragraph()
p.add_run('In Generalizer Agent: ').bold = True
p.add_run('Generate diverse design alternatives. Input: Base FBSL + strategy. Output: Modified FBSL variants.')

p = doc.add_paragraph()
p.add_run('In Research Agent: ').bold = True
p.add_run('Formulate semantic queries from design context. Input: Current FBSL. Output: Natural language query for embedding.')

p = doc.add_paragraph()
p.add_run('In Refinement Agent: ').bold = True
p.add_run('Suggest structure modifications for behavior gaps. Input: Unsatisfied behaviors. Output: Recommended structure additions.')

doc.add_heading('4.3 Prompt Engineering', level=2)
doc.add_paragraph('Temperature Settings:')
doc.add_paragraph('• Encoding: Low temperature (0.2) for consistent extraction')
doc.add_paragraph('• Generalization: Higher temperature (0.7) for creative alternatives')

# ============ SECTION 5: POSTGRESQL ============
doc.add_heading('5. POSTGRESQL DATABASE', level=1)

doc.add_heading('5.1 Purpose', level=2)
doc.add_paragraph('• Persist all design states for traceability')
doc.add_paragraph('• Enable design evolution tracking')
doc.add_paragraph('• Store evaluations for analysis')
doc.add_paragraph('• Support project management')

doc.add_heading('5.2 Schema (3 Core Tables)', level=2)

p = doc.add_paragraph()
p.add_run('Table: projects').bold = True
doc.add_paragraph('• project_id (UUID): Primary key')
doc.add_paragraph('• project_name (VARCHAR): Human-readable name')
doc.add_paragraph('• requirements (TEXT): Original user input')
doc.add_paragraph('• context (JSONB): Additional metadata')
doc.add_paragraph('• status (ENUM): active/completed/archived')

p = doc.add_paragraph()
p.add_run('Table: fbsl_nodes').bold = True
doc.add_paragraph('• node_id (UUID): Primary key')
doc.add_paragraph('• project_id (UUID): FK → projects')
doc.add_paragraph('• parent_node_id (UUID): FK → self (tree structure)')
doc.add_paragraph('• node_type (ENUM): problem/alternative/refined/aggregated')
doc.add_paragraph('• functions, behaviors, structures, layout (JSONB): FBSL components')
doc.add_paragraph('• composite_score, functional_score, behavioral_score, structural_score, layout_score, sustainability_score (FLOAT)')
doc.add_paragraph('• generation_level (INT): Depth in GoT')

p = doc.add_paragraph()
p.add_run('Table: evaluations').bold = True
doc.add_paragraph('• evaluation_id (UUID): Primary key')
doc.add_paragraph('• node_id (UUID): FK → fbsl_nodes')
doc.add_paragraph('• score_breakdown (JSONB): Detailed per-behavior scores')
doc.add_paragraph('• strengths, weaknesses, recommendations (JSONB)')

doc.add_heading('5.3 Relationships', level=2)
doc.add_paragraph('projects (1) → (N) fbsl_nodes: One project contains many design nodes')
doc.add_paragraph('fbsl_nodes (1) → (1) evaluations: Each node has one detailed evaluation')
doc.add_paragraph('fbsl_nodes (parent) → (N) fbsl_nodes (children): Nodes form a tree via parent_node_id')

# ============ SECTION 6: DATA FLOW ============
doc.add_heading('6. DATA FLOW DIAGRAM', level=1)

doc.add_paragraph('1. USER REQUIREMENTS enter via Encoder Agent (LLM parsing)')
doc.add_paragraph('2. Research Agent queries FAISS for similar precedents')
doc.add_paragraph('3. Generalizer creates 4 diverse alternatives')
doc.add_paragraph('4. Orchestrator manages GoT exploration with complexity-based adaptation')
doc.add_paragraph('5. Scoring Agent evaluates all nodes across 5 dimensions')
doc.add_paragraph('6. Pruning Agent removes weak candidates (<70%), preserves Pareto frontier')
doc.add_paragraph('7. Aggregation Agent merges compatible high-scorers (≥90%)')
doc.add_paragraph('8. Refinement Agent iteratively improves designs (Type 1/2/3)')
doc.add_paragraph('9. Layout Agent generates spatial arrangements (force-directed + A*)')
doc.add_paragraph('10. Final ranking produces top 3-5 prototypes')
doc.add_paragraph('11. PostgreSQL stores all states for persistence and analysis')

# Save the document
output_path = r'C:\Users\sanvi\OneDrive\Desktop\layout\FBSL_KAGS_Documentation.docx'
doc.save(output_path)
print(f"Word document saved to: {output_path}")
