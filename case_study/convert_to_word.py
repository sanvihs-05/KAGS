"""
Convert Validation Report to Word Document
"""

try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False
    print("python-docx not installed. Installing...")
    import subprocess
    subprocess.check_call(['py', '-m', 'pip', 'install', 'python-docx'])
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

def markdown_to_word(md_file, docx_file):
    """Convert markdown validation report to Word document"""
    
    with open(md_file, 'r', encoding='utf-8') as f:
        content = f.read()
    
    doc = Document()
    
    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    
    lines = content.split('\n')
    i = 0
    
    while i < len(lines):
        line = lines[i].rstrip()
        
        # Skip empty lines
        if not line:
            i += 1
            continue
        
        # Heading 1 (# )
        if line.startswith('# ') and not line.startswith('## '):
            heading = line[2:]
            h = doc.add_heading(heading, level=1)
            h.alignment = WD_ALIGN_PARAGRAPH.LEFT
        
        # Heading 2 (##)
        elif line.startswith('## ') and not line.startswith('### '):
            heading = line[3:]
            doc.add_heading(heading, level=2)
        
        # Heading 3 (###)
        elif line.startswith('### ') and not line.startswith('#### '):
            heading = line[4:]
            doc.add_heading(heading, level=3)
        
        # Heading 4 (####)
        elif line.startswith('#### '):
            heading = line[5:]
            doc.add_heading(heading, level=4)
        
        # Horizontal rule
        elif line.startswith('---'):
            doc.add_paragraph('_' * 80)
        
        # Code block
        elif line.startswith('```'):
            code_lines = []
            i += 1
            while i < len(lines) and not lines[i].startswith('```'):
                code_lines.append(lines[i])
                i += 1
            
            code_para = doc.add_paragraph('\n'.join(code_lines))
            code_para.style = 'No Spacing'
            for run in code_para.runs:
                run.font.name = 'Consolas'
                run.font.size = Pt(9)
            
        # Bullet list
        elif line.startswith('- ') or line.startswith('* '):
            text = line[2:]
            # Remove markdown formatting
            text = text.replace('**', '').replace('`', '').replace('✅', '[OK]').replace('⚠️', '[!]').replace('❌', '[X]').replace('✓', '[OK]')
            doc.add_paragraph(text, style='List Bullet')
        
        # Numbered list
        elif len(line) > 2 and line[0].isdigit() and line[1] == '.':
            text = line[line.index('.')+2:]
            text = text.replace('**', '').replace('`', '').replace('✅', '[OK]').replace('⚠️', '[!]').replace('❌', '[X]').replace('✓', '[OK]')
            doc.add_paragraph(text, style='List Number')
        
        # Regular paragraph
        else:
            # Clean up markdown formatting
            text = line.replace('**', '').replace('`', '').replace('✅', '[OK]').replace('⚠️', '[!]').replace('❌', '[X]').replace('✓', '[OK]')
            if text.strip():
                doc.add_paragraph(text)
        
        i += 1
    
    # Save document
    doc.save(docx_file)
    print(f"Successfully created: {docx_file}")

if __name__ == "__main__":
    md_file = "VALIDATION_REPORT.md"
    docx_file = "VALIDATION_REPORT.docx"
    
    markdown_to_word(md_file, docx_file)
