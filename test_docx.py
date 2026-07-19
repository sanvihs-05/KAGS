import traceback, sys
try:
    from pathlib import Path
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    doc = Document()
    doc.add_heading('Test', 0)
    doc.add_paragraph('Hello world')

    out = Path(r'c:\Users\sanvi\OneDrive\Desktop\layout\docs\test_output.docx')
    out.parent.mkdir(exist_ok=True)
    doc.save(str(out))
    print(f"OK: {out} ({out.stat().st_size} bytes)")
except Exception as e:
    print(f"ERROR: {e}", file=sys.stderr)
    traceback.print_exc()
