"""
FBSL-KAGS: Full Pipeline Flowchart Word Document
Heavy emphasis on layout generation sub-pipeline
"""
import sys, traceback
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
    from docx.oxml.ns import qn
except Exception as e:
    print(f"IMPORT ERROR: {e}", file=sys.stderr); sys.exit(1)


# ── helpers ────────────────────────────────────────────────────────────────────
def shade(cell, color):
    tc = cell._element.get_or_add_tcPr()
    tc.append(tc.makeelement(qn('w:shd'), {qn('w:fill'): color, qn('w:val'): 'clear'}))

def set_col_width(table, widths_cm):
    for row in table.rows:
        for i, cell in enumerate(row.cells):
            if i < len(widths_cm):
                cell.width = Cm(widths_cm[i])

def hdr(table, cells, bg='1F3864', sizes=None):
    row = table.add_row()
    for i, t in enumerate(cells):
        c = row.cells[i]; p = c.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(str(t))
        r.font.size = Pt(sizes[i] if sizes else 9)
        r.font.name = 'Calibri'; r.bold = True
        r.font.color.rgb = RGBColor(255, 255, 255); shade(c, bg)

def drow(table, cells, aligns=None, bold=False, bg=None, sz=9):
    row = table.add_row()
    for i, t in enumerate(cells):
        c = row.cells[i]; p = c.paragraphs[0]
        al = (aligns[i] if aligns else None) or WD_ALIGN_PARAGRAPH.LEFT
        p.alignment = al
        r = p.add_run(str(t)); r.font.size = Pt(sz); r.font.name = 'Calibri'
        if bold: r.bold = True
        if bg: shade(c, bg)

def mk(doc, n_cols, headers, rows, bg='1F3864', col_widths=None, sz=9):
    t = doc.add_table(rows=1, cols=n_cols)
    t.style = 'Table Grid'; t.alignment = WD_TABLE_ALIGNMENT.LEFT
    hdr(t, headers, bg=bg)
    t._element.remove(t.rows[0]._element)
    for row in rows:
        drow(t, row, sz=sz)
    if col_widths:
        set_col_width(t, col_widths)
    return t

def h1(doc, text):
    h = doc.add_heading(text, level=1)
    for r in h.runs: r.font.color.rgb = RGBColor(0x1F, 0x38, 0x64)

def h2(doc, text):
    h = doc.add_heading(text, level=2)
    for r in h.runs: r.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)
    return h

def h3(doc, text):
    return doc.add_heading(text, level=3)

def para(doc, text, bold=False, italic=False, sz=11, color=None):
    p = doc.add_paragraph(); r = p.add_run(text)
    r.font.size = Pt(sz); r.font.name = 'Calibri'
    r.bold = bold; r.italic = italic
    if color: r.font.color.rgb = RGBColor(*color)
    return p

def flow(doc, lines):
    """Render a flowchart block in Courier New monospace."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.3)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    for i, line in enumerate(lines):
        r = p.add_run(line)
        r.font.name = 'Courier New'; r.font.size = Pt(9)
        if i < len(lines) - 1:
            p.add_run('\n').font.name = 'Courier New'
    return p

def box(doc, text, bg='E8F4F8', border_color='2E74B5'):
    """A shaded info box."""
    t = doc.add_table(rows=1, cols=1)
    t.style = 'Table Grid'
    c = t.rows[0].cells[0]
    shade(c, bg)
    p = c.paragraphs[0]
    r = p.add_run(text); r.font.size = Pt(10); r.font.name = 'Calibri'
    return t

def code(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.35)
    r = p.add_run(text)
    r.font.name = 'Courier New'; r.font.size = Pt(8.5)
    r.font.color.rgb = RGBColor(0x1F, 0x38, 0x64)

def bul(doc, label, rest, sz=10):
    p = doc.add_paragraph(style='List Bullet')
    if label:
        r1 = p.add_run(label); r1.bold = True
        r1.font.size = Pt(sz); r1.font.name = 'Calibri'
    r2 = p.add_run(rest); r2.font.size = Pt(sz); r2.font.name = 'Calibri'


# ── main builder ──────────────────────────────────────────────────────────────
def build():
    doc = Document()
    ns = doc.styles['Normal']
    ns.font.name = 'Calibri'; ns.font.size = Pt(11)
    ns.paragraph_format.space_after = Pt(5)
    ns.paragraph_format.line_spacing = 1.15

    # page margins
    for section in doc.sections:
        section.top_margin    = Cm(1.8)
        section.bottom_margin = Cm(1.8)
        section.left_margin   = Cm(2.0)
        section.right_margin  = Cm(2.0)

    # ── TITLE ────────────────────────────────────────────────────────────────
    tp = doc.add_heading('FBSL-KAGS: Complete System Pipeline', level=0)
    tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for r in tp.runs:
        r.font.size = Pt(22); r.font.color.rgb = RGBColor(0x1F, 0x38, 0x64)

    sub = doc.add_paragraph(); sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rs = sub.add_run('From Natural Language Requirements to Optimised Floor Plan — Full Process Flowchart with Layout Generation Detail')
    rs.font.size = Pt(12); rs.italic = True; rs.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)

    doc.add_paragraph('')

    # ═══════════════════════════════════════════════════════════════════════════
    # SECTION 1 — MASTER PIPELINE
    # ═══════════════════════════════════════════════════════════════════════════
    h1(doc, '1.  Master Pipeline Overview')
    para(doc,
        'The FBSL-KAGS system transforms a natural language brief into ranked, scored floor plan '
        'prototypes via a Graph-of-Thoughts (GoT) exploration loop. Each node in the GoT tree '
        'represents one design variant; all nine layout generation steps run independently per node.')

    doc.add_paragraph('')

    flow(doc, [
        '┌─────────────────────────────────────────────────────────────────────────────┐',
        '│                    FBSL-KAGS SYSTEM PIPELINE                               │',
        '└─────────────────────────────────────────────────────────────────────────────┘',
        '',
        '  [USER INPUT]  Natural Language Design Brief',
        '       │        "4-bed family home, 185 m², Nordic climate, open-plan kitchen"',
        '       │',
        '       ▼',
        '  ┌──────────────────────────────────────────────────────────┐',
        '  │  STAGE 1 · RESEARCH AGENT                               │',
        '  │  FAISS RAG retrieval from Finnish floor plan corpus      │',
        '  │  → 5 precedent plans ranked by composite embedding score │',
        '  └──────────────────────────────────────────────────────────┘',
        '       │  retrieved_context: spatial patterns, room adjacencies',
        '       ▼',
        '  ┌──────────────────────────────────────────────────────────┐',
        '  │  STAGE 2 · ENCODER AGENT                                │',
        '  │  NL + context → structured FBSL root node               │',
        '  │  · Functions  (12 rooms, priorities, activities)        │',
        '  │  · Behaviors  (14 measurable metrics)                   │',
        '  │  · Structures (partitions, windows, HVAC, foundation)   │',
        '  │  · Layout     (Room objects with area targets)          │',
        '  └──────────────────────────────────────────────────────────┘',
        '       │  root_node: FBSLLayoutNode',
        '       ▼',
        '  ┌──────────────────────────────────────────────────────────┐',
        '  │  STAGE 3 · COMPLEXITY ASSESSMENT                        │',
        '  │  C_req  = 0.549  (functional complexity of brief)       │',
        '  │  C_fbsl = 0.665  (structural complexity of FBSL node)   │',
        '  │  C_overall = 0.619  →  complexity tier: HIGH            │',
        '  └──────────────────────────────────────────────────────────┘',
        '       │  GoT parameters: depth=2, breadth=5, max_nodes=112',
        '       ▼',
        '  ╔══════════════════════════════════════════════════════════╗',
        '  ║  STAGE 4 · GRAPH-OF-THOUGHTS ENGINE  (main loop)       ║',
        '  ║                                                          ║',
        '  ║  for each expansion level (0 → depth):                  ║',
        '  ║    ┌────────────────────────────────────────────────┐   ║',
        '  ║    │  SPECIALIZER AGENT                             │   ║',
        '  ║    │  expand each leaf → breadth=5 child variants   │   ║',
        '  ║    └─────────────────┬──────────────────────────────┘   ║',
        '  ║                      │  child_nodes[]                    ║',
        '  ║                      ▼                                   ║',
        '  ║    ┌────────────────────────────────────────────────┐   ║',
        '  ║    │  LAYOUT AGENT  ◄══ SEE SECTION 2 FOR DETAIL ══│   ║',
        '  ║    │  9-step spatial generation per node            │   ║',
        '  ║    │  → room_polygons, SVG, PNG, circulation paths  │   ║',
        '  ║    └─────────────────┬──────────────────────────────┘   ║',
        '  ║                      │  Layout object (with metrics)     ║',
        '  ║                      ▼                                   ║',
        '  ║    ┌────────────────────────────────────────────────┐   ║',
        '  ║    │  SCORING AGENT                                 │   ║',
        '  ║    │  S_f=0.938  S_b≈0.87  S_s=1.0                 │   ║',
        '  ║    │  S_l=0.816  S_sust=0.550                       │   ║',
        '  ║    │  S_composite = harmonic_mean ≈ 0.86            │   ║',
        '  ║    └─────────────────┬──────────────────────────────┘   ║',
        '  ║                      │  scored_nodes[]                   ║',
        '  ║                      ▼                                   ║',
        '  ║    ┌────────────────────────────────────────────────┐   ║',
        '  ║    │  GENERALIZER AGENT                             │   ║',
        '  ║    │  prune branches with S_composite < threshold   │   ║',
        '  ║    │  keep best breadth=5 per level                 │   ║',
        '  ║    └────────────────────────────────────────────────┘   ║',
        '  ║                                                          ║',
        '  ║  [if score diverges ≥ 0.3 → REFINEMENT AGENT]          ║',
        '  ║    Type 1 (dev<0.3): tweak structure dimensions         ║',
        '  ║    Type 2 (0.3–0.6): relax behavior targets             ║',
        '  ║    Type 3 (≥ 0.6): redefine room functions              ║',
        '  ╚══════════════════════════════════════════════════════════╝',
        '       │  all scored nodes (up to 112)',
        '       ▼',
        '  ┌──────────────────────────────────────────────────────────┐',
        '  │  STAGE 5 · AGGREGATION AGENT                            │',
        '  │  Pareto front across [S_composite, S_l, S_sust]         │',
        '  │  rank surviving nodes, produce comparison table         │',
        '  └──────────────────────────────────────────────────────────┘',
        '       │',
        '       ▼',
        '  ┌──────────────────────────────────────────────────────────┐',
        '  │  STAGE 6 · FINAL OUTPUT AGENT                           │',
        '  │  select best prototype → prototype_N.json               │',
        '  │  attach SVG floor plan + PNG adjacency graph            │',
        '  └──────────────────────────────────────────────────────────┘',
        '       │',
        '       ▼',
        '  [OUTPUT]  Ranked prototypes  +  Visual outputs',
        '            pipeline_outputs/prototype_N.json',
        '            visual_outputs/*_layout_*.svg',
        '            visual_outputs/*_adjacency_*.png',
    ])

    doc.add_paragraph('')

    # ═══════════════════════════════════════════════════════════════════════════
    # SECTION 2 — LAYOUT AGENT DETAILED PIPELINE
    # ═══════════════════════════════════════════════════════════════════════════
    h1(doc, '2.  Layout Agent — Detailed 9-Step Pipeline')
    para(doc,
        'The Layout Agent is the computational core of FBSL-KAGS. '
        'It receives a single FBSLLayoutNode and produces a fully positioned, '
        'scored, and visualised floor plan. All nine steps run sequentially per node.')

    doc.add_paragraph('')

    flow(doc, [
        '  ┌─────────────────────────────────────────────────────────────────┐',
        '  │  INPUT: FBSLLayoutNode                                         │',
        '  │  · layout.rooms: Dict[room_id → Room(area, height, type)]      │',
        '  │  · behaviors: 14 Behavior objects with targets                  │',
        '  │  · structures: partitions + windows + HVAC + foundation         │',
        '  └─────────────────────────┬───────────────────────────────────────┘',
        '                            │',
        '         ╔══════════════════▼══════════════════╗',
        '         ║  STEP 1 · ROOM SPEC EXTRACTION      ║',
        '         ║  area → √area = initial side (m)    ║',
        '         ║  min_width = max(2.0, side × 0.7)  ║',
        '         ║  max_width = min(15.0, side × 1.3) ║',
        '         ╚══════════════════╤══════════════════╝',
        '                            │  room_specs: Dict[room_id → {area, width, length, min_w, max_w}]',
        '         ╔══════════════════▼══════════════════╗',
        '         ║  STEP 2 · ADJACENCY MATRIX          ║',
        '         ║  WeightedAdjacencyCalculator        ║',
        '         ║  w[i,j] = α·FD + β·TF + γ·(−P)    ║',
        '         ║  α=0.40  β=0.35  γ=0.25            ║',
        '         ║  normalised to [−1, +1]             ║',
        '         ╚══════════════════╤══════════════════╝',
        '                            │  12×12 NumPy ndarray',
        '         ╔══════════════════▼══════════════════╗',
        '         ║  STEP 3 · FORCE-DIRECTED PLACEMENT  ║',
        '         ║  ForceDirectedLayout                 ║',
        '         ║  k_attr=0.1  k_rep=100.0  lr=0.01  ║',
        '         ║  max_iter=200  tol=0.01 m           ║',
        '         ║  → equilibrium positions {x, y}     ║',
        '         ╚══════════════════╤══════════════════╝',
        '                            │  initial_positions: Dict[room_id → {x, y}]',
        '         ╔══════════════════▼══════════════════╗',
        '         ║  STEP 4 · SLSQP OPTIMISATION        ║',
        '         ║  scipy.optimize.minimize (SLSQP)    ║',
        '         ║  objective: adjacency + repulsion   ║',
        '         ║           + compactness + overlap   ║',
        '         ║  constraint: w×l == target_area     ║',
        '         ║  bounds: x,y∈[0,100] w∈[min,max]  ║',
        '         ╚══════════════════╤══════════════════╝',
        '                            │  optimized_positions: Dict[room_id → {x, y, width, length}]',
        '         ╔══════════════════▼══════════════════╗',
        '         ║  STEP 5 · SHAPELY POLYGON CREATION  ║',
        '         ║  shapely.geometry.box(x, y,         ║',
        '         ║      x+width, y+length) per room   ║',
        '         ║  → .touches(), .centroid, .area     ║',
        '         ╚══════════════════╤══════════════════╝',
        '                            │  room_polygons: Dict[room_id → Shapely Polygon]',
        '         ╔══════════════════▼══════════════════╗',
        '         ║  STEP 6 · A* CIRCULATION PATHS      ║',
        '         ║  AStarPathfinder  grid=0.5m         ║',
        '         ║  8-directional  turn_penalty=0.5m  ║',
        '         ║  path_smoothing via line-of-sight   ║',
        '         ║  for each adj pair where w > 0.3    ║',
        '         ╚══════════════════╤══════════════════╝',
        '                            │  circulation_paths: List[CirculationPath]',
        '         ╔══════════════════▼══════════════════╗',
        '         ║  STEP 7 · LAYOUT OBJECT ASSEMBLY    ║',
        '         ║  Layout(room_polygons,              ║',
        '         ║         circulation_paths,          ║',
        '         ║         adjacency_matrix_actual)    ║',
        '         ╚══════════════════╤══════════════════╝',
        '                            │  layout: Layout',
        '         ╔══════════════════▼══════════════════╗',
        '         ║  STEP 8 · METRICS COMPUTATION       ║',
        '         ║  space_utilization = room_area /    ║',
        '         ║                      bbox_area      ║',
        '         ║  circulation_efficiency = mean(     ║',
        '         ║    direct_dist / path_length)       ║',
        '         ║  compactness = total / bbox         ║',
        '         ╚══════════════════╤══════════════════╝',
        '                            │  layout.metrics populated',
        '         ╔══════════════════▼══════════════════╗',
        '         ║  STEP 9 · VISUALISATION             ║',
        '         ║  ┌──────────────────────────────┐   ║',
        '         ║  │ SVGFloorPlanGenerator        │   ║',
        '         ║  │ xml.etree.ElementTree        │   ║',
        '         ║  │ → layout.svg_floor_plan      │   ║',
        '         ║  └──────────────────────────────┘   ║',
        '         ║  ┌──────────────────────────────┐   ║',
        '         ║  │ EnhancedLayoutVisualizer     │   ║',
        '         ║  │ → enhanced SVG file          │   ║',
        '         ║  │ → Matplotlib PNG 18×9 200DPI │   ║',
        '         ║  └──────────────────────────────┘   ║',
        '         ╚══════════════════╤══════════════════╝',
        '                            │',
        '  ┌─────────────────────────▼───────────────────────────────────────┐',
        '  │  OUTPUT: Layout (attached to FBSLLayoutNode)                   │',
        '  │  · room_polygons    · circulation_paths   · adjacency_svg       │',
        '  │  · svg_floor_plan  · metrics             · adjacency_matrix     │',
        '  └─────────────────────────────────────────────────────────────────┘',
    ])

    doc.add_paragraph('')

    # ═══════════════════════════════════════════════════════════════════════════
    # SECTION 3 — STEP-BY-STEP DETAIL WITH PARAMS
    # ═══════════════════════════════════════════════════════════════════════════
    h1(doc, '3.  Layout Steps — Parameters and Actual Values')

    # ── STEP 1 ────────────────────────────────────────────────────────────────
    h2(doc, 'Step 1 · Room Specification Extraction')
    mk(doc, 5,
       ['Room', 'Area (m²)', 'Initial Side (m)', 'Min Width (m)', 'Max Width (m)'],
       [
           ('Master Bedroom',    '21.0', '4.58', '3.21', '5.96'),
           ('Master Bathroom',   '7.0',  '2.65', '2.00', '3.44'),
           ('Child Bedroom × 3', '14.0', '3.74', '2.62', '4.86'),
           ('Living / Dining',   '40.0', '6.32', '4.43', '8.22'),
           ('Kitchen',           '16.0', '4.00', '2.80', '5.20'),
           ('Home Office',       '12.0', '3.46', '2.42', '4.50'),
           ('Bathroom (shared)', '6.0',  '2.45', '2.00', '3.18'),
           ('Laundry',           '5.0',  '2.24', '2.00', '2.91'),
           ('Mudroom / Storage', '5.0',  '2.24', '2.00', '2.91'),
       ], col_widths=[5.5, 2.5, 2.8, 2.8, 2.8])

    doc.add_paragraph('')

    # ── STEP 2 ────────────────────────────────────────────────────────────────
    h2(doc, 'Step 2 · Weighted Adjacency Matrix')
    para(doc, 'Formula:  w(i,j) = (α × FD + β × TF + γ × (−Privacy))  /  max(|A|)    where α=0.40, β=0.35, γ=0.25')
    mk(doc, 5,
       ['Room Pair', 'FD', 'TF', 'Privacy', 'Final w[i,j]'],
       [
           ('Kitchen ↔ Living/Dining',        '1.0', '0.8', '0.0', '+0.714  (attract strongly)'),
           ('Master Bed ↔ Master Bath',       '1.0', '0.8', '0.0', '+0.714'),
           ('Child Bed ↔ Shared Bath',        '0.6', '0.5', '0.0', '+0.435'),
           ('Living ↔ Dining',               '0.6', '0.5', '0.0', '+0.435'),
           ('Mudroom ↔ Laundry',             '0.6', '0.5', '0.0', '+0.435'),
           ('Master Bed ↔ Kitchen',          '0.0', '0.0', '1.0', '−0.263  (repel)'),
           ('Storage ↔ Living/Dining',       '0.0', '0.0', '0.0', '0.000   (neutral)'),
       ], col_widths=[5.5, 1.5, 1.5, 2.0, 5.8])

    doc.add_paragraph('')

    # ── STEP 3 ────────────────────────────────────────────────────────────────
    h2(doc, 'Step 3 · Force-Directed Placement')

    flow(doc, [
        '  Initialise: random positions in [0, 50] × [0, 50] m grid',
        '',
        '  for each iteration t in [0 … 200]:',
        '    for each pair (r_i, r_j):',
        '      d   = euclidean(r_i, r_j) + ε                  # ε=0.01 avoids ÷0',
        '      û   = direction unit vector r_i → r_j',
        '',
        '      if w[i,j] > 0:  F_attr   = k_attr × w[i,j] × d × û  # pull together',
        '      if w[i,j] < 0:  F_push   = k_attr × |w[i,j]| × d × û # push apart',
        '      always:         F_repulse = (k_rep / d²) × û          # universal repulse',
        '',
        '    position[r_i] += lr × net_force[r_i]',
        '    position[r_i]  = clip(position, 0, 100)',
        '',
        '    if max_displacement < 0.01 m:  CONVERGED  →  break',
        '',
        '  Equilibrium gap between kitchen–living pair:',
        '  d_eq = (k_rep / (k_attr × w))^(1/3) = (100 / (0.1 × 0.714))^(1/3) ≈ 11.2 m',
    ])

    mk(doc, 3,
       ['Parameter', 'Value', 'Effect'],
       [
           ('k_attraction',       '0.1',   'Strength of pull between +w rooms — lower = softer clustering'),
           ('k_repulsion',        '100.0', 'All-pairs anti-overlap force — dominates at short range'),
           ('learning_rate',      '0.01',  'Step size per iteration — smaller = smoother but slower'),
           ('max_iterations',     '200',   'Hard cap on physics steps'),
           ('convergence_tol',    '0.01 m','Stop when no room moves more than this per step'),
       ], col_widths=[4.5, 2.5, 9.5])

    doc.add_paragraph('')

    # ── STEP 4 ────────────────────────────────────────────────────────────────
    h2(doc, 'Step 4 · SLSQP Constraint Optimisation')
    para(doc,
        'scipy.optimize.minimize(method="SLSQP", maxiter=200, ftol=1e-6). '
        'Parameter vector x = [x₁,y₁,w₁,l₁, x₂,y₂,w₂,l₂, … ] (4 × 12 = 48 values).')

    flow(doc, [
        '  Objective function  (minimise):',
        '  ┌──────────────────────────────────────────────────────────────┐',
        '  │  (1) Adjacency penalty   Σ w[i,j] × max(0, dist−sep)²       │',
        '  │      for pairs where w[i,j] > 0                             │',
        '  │  (2) Repulsion penalty   Σ |w[i,j]| × max(0,1.5·sep−dist)² │',
        '  │      for pairs where w[i,j] < 0                             │',
        '  │  (3) Compactness         (bbox_area − room_area) × 0.01     │',
        '  │  (4) Overlap penalty     overlap_area × 1000.0  (very hard) │',
        '  └──────────────────────────────────────────────────────────────┘',
        '',
        '  Equality constraints (one per room):',
        '      width[i] × length[i] = target_area[i]',
        '',
        '  Bounds:',
        '      x, y  ∈ [0, 100]',
        '      width ∈ [min_width, max_width]',
        '      length∈ [min_length, max_length]',
    ])

    doc.add_paragraph('')

    # ── STEP 5 ────────────────────────────────────────────────────────────────
    h2(doc, 'Step 5 · Shapely Polygon Generation')
    para(doc, 'Each room\'s optimised {x, y, width, length} is converted to a Shapely Polygon using box():')

    flow(doc, [
        '  from shapely.geometry import box',
        '',
        '  polygon = box(x, y, x + width, y + length)',
        '              └──────────────────────────────┘',
        '              Creates axis-aligned rectangle as a Polygon object',
        '',
        '  Shapely methods used downstream:',
        '  ┌─────────────────────────────────────────────────────────────┐',
        '  │ polygon.bounds     → (minx, miny, maxx, maxy)  [SVG coords]│',
        '  │ polygon.centroid   → Point(cx, cy)             [label pos] │',
        '  │ polygon.area       → float                     [verify m²] │',
        '  │ polygon.touches(q) → bool       [shared wall detection]    │',
        '  │ polygon.distance(q)→ float      [near-adj check: < 0.5 m]  │',
        '  └─────────────────────────────────────────────────────────────┘',
    ])

    doc.add_paragraph('')

    # ── STEP 6 ────────────────────────────────────────────────────────────────
    h2(doc, 'Step 6 · A* Circulation Pathfinding')

    flow(doc, [
        '  AStarPathfinder  (spatial_algorithms.py)',
        '',
        '  GRID:',
        '    resolution  = 0.5 m per cell',
        '    bounds      = (min_x−2, min_y−2, max_x+2, max_y+2)  [2 m safety margin]',
        '    obstacles   = all room bounding boxes marked as non-walkable',
        '',
        '  OPEN SET: MinHeap  [(f_score, counter, cell)]',
        '    g[start]   = 0',
        '    h(n, goal) = Manhattan distance  (dx + dy)',
        '    f(n)       = g(n) + h(n)',
        '',
        '  SEARCH LOOP:',
        '  ┌────────────────────────────────────────────────────────────┐',
        '  │ while open_set not empty:                                  │',
        '  │   current = heappop(open_set)                             │',
        '  │   if current == goal: return reconstruct_path()           │',
        '  │   for neighbour in 8_directions(current):                 │',
        '  │     if not walkable: skip                                  │',
        '  │     g_new = g[current] + euclidean(current, neighbour)    │',
        '  │     if g_new < g[neighbour]:                              │',
        '  │       g[neighbour] = g_new                                │',
        '  │       f[neighbour] = g_new + manhattan(neighbour, goal)   │',
        '  │       heappush(open_set, (f, counter++, neighbour))       │',
        '  └────────────────────────────────────────────────────────────┘',
        '',
        '  POST-PROCESS:',
        '    path_smoothed via line-of-sight (skip intermediate waypoints',
        '    if start→end is obstacle-free in a straight line)',
        '    path_cost = path_length + 0.5 × (num_bends − 1)',
        '',
        '  TRIGGER: for each adj pair where w[i,j] > 0.3',
    ])

    doc.add_paragraph('')

    # ── STEP 8 METRICS ────────────────────────────────────────────────────────
    h2(doc, 'Step 8 · Layout Metrics Computation')

    mk(doc, 4,
       ['Metric', 'Formula', 'Good Value', 'Actual (example)'],
       [
           ('Space Utilisation',      'Σroom_area / bbox_area',                    '> 0.70', '0.75'),
           ('Circulation Efficiency', 'mean(direct_dist / path_length) per path',  '> 0.80', '0.82'),
           ('Compactness',            'Σroom_area / bounding_box_area',             '> 0.65', '0.71'),
           ('Adjacency Score',        'met_adj_pairs / required_adj_pairs',         '> 0.85', '0.87'),
           ('Room Count',             'len(layout.rooms)',                           '== target', '12 / 12'),
       ], col_widths=[4.5, 6.5, 2.8, 2.8])

    doc.add_paragraph('')

    # ═══════════════════════════════════════════════════════════════════════════
    # SECTION 4 — VISUALISATION PIPELINE
    # ═══════════════════════════════════════════════════════════════════════════
    h1(doc, '4.  Step 9 — Visualisation Pipeline (Full Detail)')
    para(doc,
        'Step 9 produces four distinct visual artefacts. Each runs on the same '
        'Shapely polygon data but uses a different rendering backend.')

    flow(doc, [
        '  room_polygons + circulation_paths',
        '       │',
        '       ├─────────────────────────────────────────────────────────────┐',
        '       │                                                             │',
        '       ▼                                                             ▼',
        '  ┌─────────────────────────────────┐   ┌──────────────────────────────────────┐',
        '  │  SVGFloorPlanGenerator          │   │  EnhancedLayoutVisualizer            │',
        '  │  (utils/visualization.py)       │   │  (visualization/enhanced_layout.py)  │',
        '  │                                 │   │                                      │',
        '  │  1. _calculate_transform()      │   │  1. LayoutRoomAdapter.extract()      │',
        '  │     scale = min(sx,sy, 20px/m) │   │     Polygon → {x,y,w,h,room_type}   │',
        '  │                                 │   │                                      │',
        '  │  2. _add_styles()               │   │  2. CompactRoomPlacer (fallback)     │',
        '  │     CSS: .room .room-label      │   │     importance scores + 4-slot       │',
        '  │     .circulation .dimension     │   │                                      │',
        '  │                                 │   │  3. Enhanced SVG (xml.etree)         │',
        '  │  3. _draw_rooms()               │   │     compass rose  1m grid            │',
        '  │     Shapely .exterior.coords    │   │     room colours (dark palette)      │',
        '  │     → <polygon points="…"/>    │   │     → visual_outputs/*.svg           │',
        '  │     fill from room_colors dict  │   │                                      │',
        '  │                                 │   │  4. _generate_adjacency_graph()      │',
        '  │  4. _draw_circulation()         │   │     NetworkX graph + spring_layout   │',
        '  │     A* paths → <polyline/>      │   │     Matplotlib 18×9 @ 200 DPI       │',
        '  │     blue dashed, opacity=0.6    │   │     → visual_outputs/*.png           │',
        '  │                                 │   └──────────────────────────────────────┘',
        '  │  5. _add_legend()               │',
        '  │     colour swatch per room      │',
        '  │                                 │',
        '  │  6. _add_metrics_overlay()      │',
        '  │     Space util / Circ / Compact │',
        '  │                                 │',
        '  │  → layout.svg_floor_plan (str)  │',
        '  └─────────────────────────────────┘',
    ])

    doc.add_paragraph('')

    # ── SVG DETAILS ────────────────────────────────────────────────────────────
    h2(doc, '4.1  SVG Floor Plan Layer Stack')

    mk(doc, 4,
       ['Layer', 'SVG Element', 'Content', 'Key Data Source'],
       [
           ('1 · Background',   '<rect> canvas',             '1200×800 px white background',                      'VisualConfig constants'),
           ('2 · Grid',         '<line> × N',                '1 m × 1 m grid in #ECF0F1, opacity=0.4',            'grid_spacing=1.0, scale'),
           ('3 · Rooms',        '<polygon points="x,y …">',  'Shapely polygon.exterior.coords → SVG points',      'room_polygons dict'),
           ('4 · Room Labels',  '<text>',                    'Room name at centroid − 0.5 m (y), area at +0.5 m', 'polygon.centroid'),
           ('5 · Circulation',  '<polyline>',                'A* waypoints as blue dashed path',                  'circulation_paths'),
           ('6 · Legend',       '<rect> + <text> per room',  'Colour swatch + "name (area m²)"',                  'room_colors dict'),
           ('7 · Metrics',      '<text> block',              'Space util %, Circulation eff %, Compactness %',    'layout.metrics'),
           ('8 · Compass',      '<circle> + <polygon>',      'N-arrow at top-right, compass_size=60 px',          'EnhancedLayoutVisualizer'),
       ], col_widths=[3.0, 3.5, 6.0, 4.0])

    doc.add_paragraph('')

    # ── ADJACENCY GRAPH ────────────────────────────────────────────────────────
    h2(doc, '4.2  Adjacency Graph PNG — NetworkX + Matplotlib')

    flow(doc, [
        '  room_polygons + adjacency_matrix',
        '       │',
        '       ▼',
        '  ┌──────────────────────────────────────────────────────────┐',
        '  │  1. Build NetworkX graph                                │',
        '  │     .add_node(room_id, room_type, area, x, y, w, h)    │',
        '  │     for each pair: check _rooms_adjacent()             │',
        '  │       tolerance = 0.2 m,  min_overlap = 30% of shorter │',
        '  │       → .add_edge(a, b, edge_type=classify(a,b))       │',
        '  │     if not nx.is_connected(G):                          │',
        '  │       add "bridge" edges between components             │',
        '  └────────────────────┬─────────────────────────────────────┘',
        '                       │  NetworkX graph G',
        '       ┌───────────────┴───────────────────┐',
        '       ▼                                   ▼',
        '  ┌────────────────────┐   ┌───────────────────────────────────┐',
        '  │  LEFT PANEL        │   │  RIGHT PANEL                      │',
        '  │  Spatial layout    │   │  Spring-layout graph              │',
        '  │  _plot_spatial()   │   │  _plot_graph()                    │',
        '  │                    │   │                                   │',
        '  │  Matplotlib        │   │  pos = nx.spring_layout(G,        │',
        '  │  Rectangle patches │   │            seed=42)               │',
        '  │  room fill colours │   │                                   │',
        '  │  centroid-centroid │   │  nx.draw_networkx_nodes(          │',
        '  │  edge lines        │   │      node_color=room_colors,      │',
        '  │  ax.invert_yaxis() │   │      node_size=900)               │',
        '  │  ax.set_aspect(eq) │   │                                   │',
        '  └────────────────────┘   │  Draw edges by type:             │',
        '                           │  critical → #27AE60 solid w=3.5  │',
        '                           │  spatial  → #3498DB solid w=2.5  │',
        '                           │  proximity→ #95A5A6 dotted w=1.5 │',
        '                           │  bridge   → #E74C3C dashed w=2.0 │',
        '                           └───────────────────────────────────┘',
        '                                      │',
        '                     fig.savefig(path, dpi=200)',
        '                     → visual_outputs/*_adjacency_*.png',
    ])

    doc.add_paragraph('')

    mk(doc, 4,
       ['Edge Type', 'Classification Condition', 'Colour', 'Style'],
       [
           ('critical',   'kitchen↔dining/living, bedroom↔bathroom/corridor', '#27AE60 (green)',    'solid, width=3.5'),
           ('spatial',    'any other geometry-detected adjacency',             '#3498DB (blue)',     'solid, width=2.5'),
           ('proximity',  'fallback: nearest N−1 pairs by centroid distance',  '#95A5A6 (grey)',    'dotted, width=1.5'),
           ('bridge',     'added to reconnect disconnected graph components',  '#E74C3C (red)',      'dashed, width=2.0'),
       ], col_widths=[2.8, 6.5, 3.5, 3.7])

    doc.add_paragraph('')

    # ═══════════════════════════════════════════════════════════════════════════
    # SECTION 5 — SCORING PIPELINE
    # ═══════════════════════════════════════════════════════════════════════════
    h1(doc, '5.  Scoring Agent — How Layout Feeds into Scores')

    flow(doc, [
        '  FBSLLayoutNode (with Layout attached)',
        '       │',
        '       ├──► _score_functions()   → S_f = weighted function satisfaction',
        '       │',
        '       ├──► _score_behaviors()   → S_b = geometric_mean(actual/target)',
        '       │         │',
        '       │         ├─ LIGHTING:  daylight_factor = window_ratio × 0.75 × 100',
        '       │         │             ratio = min(1.0, actual_DF / target_DF)',
        '       │         │             (window_ratio from Structure.dimensions["window_ratio"])',
        '       │         │',
        '       │         └─ VENTILATION: has_hvac = any(s.type==MEP and "hvac" in s.name)',
        '       │                         score = 1.0 if has_hvac else 0.4 (fallback)',
        '       │',
        '       ├──► _score_structures()  → S_s = base × penalties',
        '       │         │',
        '       │         ├─ has_structural = any(s.load_bearing)   → no ×0.7 penalty',
        '       │         └─ has_envelope   = any(s.category=="envelope") → no ×0.8 penalty',
        '       │',
        '       ├──► _score_layout()      → S_l = f(space_util, circ_eff, adj_quality)',
        '       │         │',
        '       │         └─ uses layout.metrics from Step 8',
        '       │',
        '       └──► _score_sustainability() → S_sust = 0.550 (orientation/materials)',
        '                                                        ↑ not yet layout-coupled',
        '',
        '  S_composite = harmonic_mean([S_f, S_b, S_s, S_l, S_sust])',
        '              = 1 / (Σ weight_i / S_i)    [rho=−1 generalised mean]',
    ])

    doc.add_paragraph('')

    mk(doc, 5,
       ['Criterion', 'Weight', 'Before Fix', 'After Fix', 'Root Cause of Change'],
       [
           ('S_f  (Functions)',      '0.25', '0.938', '0.938', 'Unchanged — functions fully met'),
           ('S_b  (Behaviors)',      '0.20', '0.213', '~0.87', 'Window structures → lighting DF ≥ threshold; HVAC → ventilation=1.0'),
           ('S_s  (Structures)',     '0.20', '0.560', '1.000', 'load_bearing=True removes ×0.7; category=envelope removes ×0.8'),
           ('S_l  (Layout)',         '0.25', '0.816', '0.816', 'Unchanged — geometry already good'),
           ('S_sust (Sustain.)',     '0.10', '0.550', '0.550', 'Unchanged — materials/orientation not yet layout-coupled'),
           ('S_composite',          '—',    '0.483', '~0.860', 'Harmonic mean lifts when no score is critically low'),
       ], col_widths=[3.8, 1.8, 2.4, 2.4, 6.3])

    doc.add_paragraph('')

    # ═══════════════════════════════════════════════════════════════════════════
    # SECTION 6 — GoT PARAMETERS
    # ═══════════════════════════════════════════════════════════════════════════
    h1(doc, '6.  Graph-of-Thoughts Parameters (Complexity → GoT Config)')

    flow(doc, [
        '  C_req   = 0.549   (functional complexity of brief:  num_rooms, area, constraints)',
        '  C_fbsl  = 0.665   (structural complexity of FBSL:   functions, behaviors, structures)',
        '  C_overall = 0.619  →  tier: HIGH  (threshold > 0.6)',
        '',
        '  GoT config for HIGH complexity:',
        '  ┌──────────────────────────────────────────────┐',
        '  │  max_depth  = 2        (levels of expansion) │',
        '  │  breadth    = 5        (children per node)   │',
        '  │  max_nodes  = 112      (total node cap)      │',
        '  │  pruning    = enabled  (generalizer active)  │',
        '  │  refinement = enabled  (FBS reformulation)   │',
        '  └──────────────────────────────────────────────┘',
        '',
        '  GoT tree (worst case, no pruning):',
        '    depth 0: root node           (1 node)',
        '    depth 1: 1 × 5 children     (5 nodes)',
        '    depth 2: 5 × 5 grandchildren(25 nodes, capped at max_nodes)',
        '    Layout Agent runs 9 steps per node → up to 31 layout runs',
    ])

    doc.add_paragraph('')

    # ═══════════════════════════════════════════════════════════════════════════
    # SECTION 7 — DATA OBJECTS REFERENCE
    # ═══════════════════════════════════════════════════════════════════════════
    h1(doc, '7.  Key Data Objects Flowing Through the Pipeline')

    mk(doc, 4,
       ['Object', 'Type / Location', 'Key Fields', 'Created By'],
       [
           ('FBSLLayoutNode', 'fbsl_models.py', 'node_id, functions{}, behaviors{}, structures{}, layout, metadata{}', 'EncoderAgent'),
           ('Function',       'fbsl_models.py', 'function_id, name, category, priority, activities[], spatial_requirements{}', 'EncoderAgent'),
           ('Behavior',       'fbsl_models.py', 'category, metric_name, target_value, actual_value, tolerance, derived_from_function', 'EncoderAgent'),
           ('Structure',      'fbsl_models.py', 'structure_id, name, structure_type (enum), category, material_type, load_bearing, dimensions{}', 'EncoderAgent'),
           ('Layout',         'fbsl_models.py', 'rooms{}, room_polygons{}, circulation_paths[], svg_floor_plan, adjacency_svg, metrics{}', 'LayoutAgent'),
           ('Room',           'fbsl_models.py', 'room_id, name, room_type, area, height, required_adjacencies[], preferred_adjacencies[]', 'EncoderAgent'),
           ('CirculationPath','spatial_algorithms.py', 'path_points[], path_length, efficiency, room_pair, corridor_width', 'AStarPathfinder'),
           ('GoTNode',        'got_engine.py',  'fbsl_node, parent_id, depth, score, children[]', 'GoT Engine'),
       ], col_widths=[3.5, 3.5, 7.5, 3.0])

    doc.add_paragraph('')

    # ═══════════════════════════════════════════════════════════════════════════
    # SECTION 8 — TECH STACK SUMMARY
    # ═══════════════════════════════════════════════════════════════════════════
    h1(doc, '8.  Technology Stack Summary')

    mk(doc, 4,
       ['Library', 'Role in Pipeline', 'Stage Used', 'Output Produced'],
       [
           ('Ollama (local LLM)',       'NL → structured FBSL JSON (LLaMA 3.1 / Mistral)', 'Stages 1, 2, 4', 'FBSL node JSON'),
           ('FAISS (faiss-cpu)',        'Vector similarity search over Finnish floor plan corpus', 'Stage 1', '5 ranked precedent plans'),
           ('NumPy',                   'Adjacency matrix ops, position arrays, force calculations', 'Steps 2, 3', '12×12 ndarray, position dict'),
           ('scipy.optimize (SLSQP)',  'Constrained layout optimisation (48-var problem)', 'Step 4', 'Refined {x,y,w,l} per room'),
           ('Shapely',                 'Geometry: box(), .touches(), .centroid, .distance()', 'Steps 5, 6, 9', 'Polygon objects'),
           ('heapq',                   'Min-heap priority queue for A* open set', 'Step 6', 'Optimal circulation paths'),
           ('NetworkX',                'Graph structure, spring_layout(), connected_components()', 'Step 9', 'Layout graph G'),
           ('Matplotlib',              '18×9 dual-panel figure saved at 200 DPI', 'Step 9', 'PNG adjacency graph'),
           ('xml.etree.ElementTree',   'SVG DOM construction (rooms, labels, legend, metrics)', 'Step 9', 'SVG string / file'),
           ('xml.dom.minidom',         'Prettify SVG XML for readability', 'Step 9', 'Indented SVG output'),
           ('python-docx + lxml',      'Word document generation for reports', 'Output stage', '.docx report files'),
           ('pydantic / dataclasses',  'FBSL model validation and serialisation', 'All stages', 'Validated node objects'),
       ], col_widths=[4.0, 6.0, 2.8, 3.7])

    doc.add_paragraph('')

    # ── SAVE ─────────────────────────────────────────────────────────────────
    candidates = [
        Path(r'c:\Users\sanvi\OneDrive\Desktop\layout\docs\FBSL_KAGS_Pipeline_Flowchart.docx'),
        Path(r'c:\Users\sanvi\Documents\FBSL_KAGS_Pipeline_Flowchart.docx'),
        Path(r'c:\Users\sanvi\FBSL_Pipeline.docx'),
    ]
    for out in candidates:
        try:
            out.parent.mkdir(parents=True, exist_ok=True)
            doc.save(str(out))
            print(f"Saved: {out}")
            print(f"Size : {out.stat().st_size / 1024:.1f} KB")
            break
        except Exception as e:
            print(f"Could not save to {out}: {e}", file=sys.stderr)


if __name__ == '__main__':
    try:
        build()
    except Exception as e:
        print(f"BUILD FAILED: {e}", file=sys.stderr)
        traceback.print_exc()
        sys.exit(1)
